from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/Volumes/Users/Public/work/基于大数据的中国金融期货交易分析系统")
SRC = ROOT / "1_修改完善版.docx"
OUT = ROOT / "1_规范排版版.docx"

CN_TITLE = "中国金融期货大数据分析与风险评估系统设计与实现"
CN_TITLE_BREAK = "中国金融期货大数据分析与风险评估系统\n设计与实现"
EN_TITLE_BREAK = (
    "Design and Implementation of a Big Data Analytics and\n"
    "Risk Assessment System for Chinese Financial Futures"
)
EN_TITLE_FORMATTED = (
    "Design and Implementation of a Big Data Analytics\n"
    "and Risk Assessment System for Chinese Financial Futures"
)


def clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_run_font(
    run,
    east_asia: str,
    ascii_font: str,
    size_pt: float,
    bold: bool | None = None,
) -> None:
    run.font.name = ascii_font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def style_single_run_paragraph(
    paragraph,
    *,
    east_asia: str,
    ascii_font: str,
    size_pt: float,
    bold: bool,
    align: WD_ALIGN_PARAGRAPH | None,
    line_spacing_pt: float | None = None,
    space_before_pt: float | None = None,
    space_after_pt: float | None = None,
    first_line_indent_cm: float | None = None,
) -> None:
    if align is not None:
        paragraph.alignment = align
    pf = paragraph.paragraph_format
    if line_spacing_pt is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing_pt)
    if space_before_pt is not None:
        pf.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        pf.space_after = Pt(space_after_pt)
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    elif first_line_indent_cm == 0:
        pf.first_line_indent = Cm(0)

    for run in paragraph.runs:
        if run.text:
            set_run_font(run, east_asia, ascii_font, size_pt, bold)


def keep_at_most_n_blank_between(doc: Document, start_idx: int, end_idx: int, keep: int) -> None:
    blanks_seen = 0
    for paragraph in list(doc.paragraphs[start_idx + 1 : end_idx]):
        if paragraph.text.strip():
            blanks_seen = 0
            continue
        if paragraph._p.pPr is not None and paragraph._p.pPr.sectPr is not None:
            blanks_seen = 0
            continue
        if any(
            br.get(qn("w:type")) == "page"
            for br in paragraph._p.findall(".//" + qn("w:br"))
        ):
            blanks_seen = 0
            continue
        blanks_seen += 1
        if blanks_seen > keep:
            delete_paragraph(paragraph)


def add_page_field(paragraph, literal: str | None = None) -> None:
    clear_runs(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if literal is not None:
        run = paragraph.add_run(literal)
        set_run_font(run, "Times New Roman", "Times New Roman", 10.5, False)
        return

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
    set_run_font(run, "Times New Roman", "Times New Roman", 10.5, False)


def set_section_page_numbering(section, fmt: str, start: int | None = None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))
    elif pg_num_type.get(qn("w:start")) is not None:
        del pg_num_type.attrib[qn("w:start")]


def set_header_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def style_header_text(paragraph, text: str) -> None:
    clear_runs(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", 10.5, False)
    set_header_border(paragraph)


def set_style_font(style, east_asia: str, ascii_font: str, size_pt: float, bold: bool) -> None:
    style.font.name = ascii_font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.element.rPr.get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)


def is_reference_paragraph(text: str) -> bool:
    stripped = text.strip()
    return stripped.startswith("[") and "]" in stripped[:6]


doc = Document(SRC)
doc.core_properties.title = CN_TITLE
doc.core_properties.author = "郭子健"
doc.core_properties.subject = "上海立达学院本科毕业论文"

# Global page setup.
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

doc.settings.odd_and_even_pages_header_footer = True

# Remove the extra blank page between Chinese abstract and English abstract.
kw_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("关键词："))
eng_abs_marker_idx = next(
    i
    for i, p in enumerate(doc.paragraphs)
    if i > kw_idx and "BACHELOR'S DEGREE THESIS OF SHANGHAI LIDA UNIVERSITY" in p.text
)
for paragraph in list(doc.paragraphs[kw_idx + 1 : eng_abs_marker_idx]):
    has_page_break = any(
        br.get(qn("w:type")) == "page" for br in paragraph._p.findall(".//" + qn("w:br"))
    )
    if (
        not paragraph.text.strip()
        and not has_page_break
        and not (paragraph._p.pPr is not None and paragraph._p.pPr.sectPr is not None)
    ):
        delete_paragraph(paragraph)
eng_abs_marker = next(
    p
    for i, p in enumerate(doc.paragraphs)
    if i > kw_idx and "BACHELOR'S DEGREE THESIS OF SHANGHAI LIDA UNIVERSITY" in p.text
)
clear_runs(eng_abs_marker)
eng_abs_marker.add_run(EN_TITLE_FORMATTED)

# Tighten English inner cover and section transitions to avoid blank spill pages.
eng_cover_idx = next(
    i
    for i, p in enumerate(doc.paragraphs)
    if "BACHELOR'S DEGREE THESIS OF SHANGHAI LIDA UNIVERSITY" in p.text
)
eng_title_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == EN_TITLE_BREAK)
candidate_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Candidate：Guo Zijian")
keep_at_most_n_blank_between(doc, eng_cover_idx, eng_title_idx, keep=1)
eng_title_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == EN_TITLE_BREAK)
candidate_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Candidate：Guo Zijian")
date_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Mar. 20th, 2026")
keep_at_most_n_blank_between(doc, eng_title_idx, candidate_idx, keep=1)
candidate_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Candidate：Guo Zijian")
date_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Mar. 20th, 2026")
keep_at_most_n_blank_between(doc, candidate_idx, date_idx, keep=1)

date_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Mar. 20th, 2026")
declaration_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "郑重声明")
for paragraph in list(doc.paragraphs[date_idx + 1 : declaration_idx]):
    if (
        not paragraph.text.strip()
        and not (paragraph._p.pPr is not None and paragraph._p.pPr.sectPr is not None)
        and not any(
            br.get(qn("w:type")) == "page"
            for br in paragraph._p.findall(".//" + qn("w:br"))
        )
    ):
        delete_paragraph(paragraph)
next(p for p in doc.paragraphs if p.text.strip() == "郑重声明").paragraph_format.page_break_before = True

cn_abstract_title_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == CN_TITLE)
last_decl_content_idx = max(
    i for i, p in enumerate(doc.paragraphs[:cn_abstract_title_idx]) if p.text.strip()
)
for paragraph in list(doc.paragraphs[last_decl_content_idx + 1 : cn_abstract_title_idx]):
    if (
        not paragraph.text.strip()
        and not (paragraph._p.pPr is not None and paragraph._p.pPr.sectPr is not None)
        and not any(
            br.get(qn("w:type")) == "page"
            for br in paragraph._p.findall(".//" + qn("w:br"))
        )
    ):
        delete_paragraph(paragraph)

# Cover title and English cover title.
for paragraph in doc.paragraphs:
    stripped = paragraph.text.strip()
    if stripped == CN_TITLE_BREAK:
        style_single_run_paragraph(
            paragraph,
            east_asia="黑体",
            ascii_font="Times New Roman",
            size_pt=22,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing_pt=32,
            first_line_indent_cm=0,
        )
    elif stripped == EN_TITLE_BREAK:
        clear_runs(paragraph)
        paragraph.add_run(EN_TITLE_FORMATTED)
        style_single_run_paragraph(
            paragraph,
            east_asia="Times New Roman",
            ascii_font="Times New Roman",
            size_pt=16,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing_pt=22,
        )
    elif stripped == EN_TITLE_FORMATTED:
        style_single_run_paragraph(
            paragraph,
            east_asia="Times New Roman",
            ascii_font="Times New Roman",
            size_pt=16,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing_pt=22,
            first_line_indent_cm=0,
        )

# Front-matter headings and keyword lines.
for paragraph in doc.paragraphs:
    stripped = paragraph.text.strip()
    if stripped == CN_TITLE:
        style_single_run_paragraph(
            paragraph,
            east_asia="黑体",
            ascii_font="Times New Roman",
            size_pt=22,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing_pt=32,
            first_line_indent_cm=0,
        )
    elif stripped == "摘要":
        style_single_run_paragraph(
            paragraph,
            east_asia="黑体",
            ascii_font="Times New Roman",
            size_pt=18,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing_pt=32,
            first_line_indent_cm=0,
        )
    elif stripped == "ABSTRACT":
        style_single_run_paragraph(
            paragraph,
            east_asia="Times New Roman",
            ascii_font="Times New Roman",
            size_pt=18,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing_pt=23,
            first_line_indent_cm=0,
        )
    elif stripped == "目　录":
        style_single_run_paragraph(
            paragraph,
            east_asia="黑体",
            ascii_font="Times New Roman",
            size_pt=18,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing_pt=32,
            first_line_indent_cm=0,
        )
    elif stripped.startswith("关键词："):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0)
    elif stripped.startswith("Key words:"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0)

# Heading styles.
heading1 = doc.styles["Heading 1"]
set_style_font(heading1, "黑体", "Times New Roman", 18, True)
heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
heading1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
heading1.paragraph_format.line_spacing = Pt(23)
heading1.paragraph_format.space_before = Pt(18)
heading1.paragraph_format.space_after = Pt(12)

heading2 = doc.styles["Heading 2"]
set_style_font(heading2, "黑体", "Times New Roman", 14, True)
heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
heading2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
heading2.paragraph_format.line_spacing = Pt(23)
heading2.paragraph_format.space_before = Pt(12)
heading2.paragraph_format.space_after = Pt(12)

if "Heading 3" in [style.name for style in doc.styles]:
    heading3 = doc.styles["Heading 3"]
    set_style_font(heading3, "黑体", "Times New Roman", 12, True)
    heading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading3.paragraph_format.line_spacing = Pt(23)

# Table of contents styles.
toc1 = doc.styles["toc 1"]
set_style_font(toc1, "黑体", "Times New Roman", 14, True)
toc1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
toc1.paragraph_format.space_after = Pt(2)

toc2 = doc.styles["toc 2"]
set_style_font(toc2, "宋体", "Times New Roman", 12, False)
toc2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
toc2.paragraph_format.space_after = Pt(2)

# Reference heading and reference paragraphs.
ref_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "参考文献")
thanks_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "致谢")
ref_heading = doc.paragraphs[ref_idx]
style_single_run_paragraph(
    ref_heading,
    east_asia="黑体",
    ascii_font="Times New Roman",
    size_pt=18,
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    line_spacing_pt=23,
    first_line_indent_cm=0,
)

for paragraph in doc.paragraphs[ref_idx + 1 : thanks_idx]:
    if not paragraph.text.strip():
        continue
    if is_reference_paragraph(paragraph.text):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = paragraph.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(23)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = Cm(0.74)
        pf.first_line_indent = Cm(-0.74)
        for run in paragraph.runs:
            if run.text:
                set_run_font(run, "宋体", "Times New Roman", 12, False)

# Front-matter and body page numbers.
sec0, sec1, sec2, sec3 = doc.sections
sec0.different_first_page_header_footer = True
for footer in [sec0.footer, sec0.even_page_footer, sec0.first_page_footer]:
    for paragraph in footer.paragraphs:
        clear_runs(paragraph)

sec1.different_first_page_header_footer = True
for paragraph in sec1.first_page_footer.paragraphs:
    clear_runs(paragraph)
for footer in [sec1.footer, sec1.even_page_footer]:
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    add_page_field(paragraph, literal="I")
set_section_page_numbering(sec1, "upperRoman")

for footer in [sec2.footer, sec2.even_page_footer]:
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    add_page_field(paragraph)
set_section_page_numbering(sec2, "upperRoman", start=2)

for paragraph in sec2.header.paragraphs:
    clear_runs(paragraph)
for paragraph in sec2.even_page_header.paragraphs:
    clear_runs(paragraph)

for footer in [sec3.footer, sec3.even_page_footer]:
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    add_page_field(paragraph)
set_section_page_numbering(sec3, "decimal", start=1)

header_text = "上海立达学院本科毕业论文（设计）"
for header in [sec3.header, sec3.even_page_header]:
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    style_header_text(paragraph, header_text)

# Keep section starts stable.
sec0.start_type = WD_SECTION_START.NEW_PAGE
sec1.start_type = WD_SECTION_START.NEW_PAGE
sec2.start_type = WD_SECTION_START.NEW_PAGE
sec3.start_type = WD_SECTION_START.NEW_PAGE

# TOC page numbers must reflect body numbering from 1.
toc_pages = {
    "1 绪论": 1,
    "1.1 研究背景与意义": 1,
    "1.2 国内外研究现状": 1,
    "1.3 研究内容与研究方法": 3,
    "1.4 论文结构安排": 3,
    "2 相关技术与理论基础": 4,
    "2.1 金融期货盘后分析的数据特征": 4,
    "2.2 系统实现相关技术": 4,
    "2.3 波动率与风险度量理论": 5,
    "2.4 本文方案的适用性说明": 6,
    "3 需求分析与数据资源设计": 7,
    "3.1 系统建设目标": 7,
    "3.2 功能需求分析": 7,
    "3.3 非功能需求分析": 8,
    "3.4 数据源设计": 8,
    "3.5 数据模型设计": 9,
    "4 系统总体设计": 10,
    "4.1 总体架构设计": 10,
    "4.2 模块划分设计": 11,
    "4.3 核心数据表设计": 11,
    "4.4 API 与页面设计": 12,
    "5 系统详细实现与关键技术": 13,
    "5.1 数据采集实现": 13,
    "5.2 主力连续序列构建": 14,
    "5.3 波动率预测与风险评估实现": 15,
    "5.4 服务层实现": 16,
    "5.5 前端展示实现": 17,
    "5.6 本章小结": 20,
    "6 系统测试与结果分析": 21,
    "6.1 测试环境与测试策略": 21,
    "6.2 功能测试结果": 21,
    "6.3 数据规模与质量分析": 22,
    "6.4 分析结果展示": 23,
    "6.5 测试与结果分析总结": 26,
    "7 结论与展望": 27,
    "7.1 研究结论": 27,
    "7.2 项目价值与不足": 27,
    "7.3 后续改进方向": 28,
    "参考文献": 29,
    "致谢": 31,
}
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith("toc"):
        title = paragraph.text.split("\t", 1)[0].strip()
        if title in toc_pages:
            clear_runs(paragraph)
            run = paragraph.add_run(f"{title}\t{toc_pages[title]}")
            if paragraph.style.name == "toc 1":
                set_run_font(run, "黑体", "Times New Roman", 14, True)
            else:
                set_run_font(run, "宋体", "Times New Roman", 12, False)

# Figure and table captions: use chapter-dot numbering and prescribed caption style.
caption_overrides = {
    "图6-6 数据质量统计图": "图6.1 数据质量统计图",
    "图6-1 四类股指期货主力连续归一化走势": "图6.2 四类股指期货主力连续归一化走势",
    "图6-2 IF 主力连续收盘价与 20 日均线": "图6.3 IF 主力连续收盘价与 20 日均线",
    "图6-3 IF 回撤与波动率结果图": "图6.4 IF 回撤与波动率结果图",
    "图6-4 各品种 60 日历史模拟 VaR 对比图": "图6.5 各品种 60 日历史模拟 VaR 对比图",
    "图6-5 股指期货与宏观指标相关性热力图": "图6.6 股指期货与宏观指标相关性热力图",
}
caption_pattern = re.compile(r"^([图表])(\d+)-(\d+)(.*)$")
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    replacement = caption_overrides.get(text)
    if replacement is None:
        match = caption_pattern.match(text)
        if match:
            replacement = f"{match.group(1)}{match.group(2)}.{match.group(3)}{match.group(4)}"
    if replacement:
        clear_runs(paragraph)
        run = paragraph.add_run(replacement)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = paragraph.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(23)
        set_run_font(run, "黑体", "Times New Roman", 12, True)

# Tighten reference punctuation and electronic-resource GB/T 7714 style.
refs = [
    "[1] 陈其安, 张慧, 陈抒妤. 股指期货交易加剧了中国股票市场波动性吗？基于投资者结构的理论和实证研究[J]. 中国管理科学, 2020, 28(4): 1-13.",
    "[2] Bohl M T, Siklos P L, Diesteldorf J. The effect of index futures trading on volatility: Three markets for Chinese stocks[J]. China Economic Review, 2015, 34: 207-224.",
    "[3] Engle R F. Autoregressive Conditional Heteroscedasticity with Estimates of the Variance of United Kingdom Inflation[J]. Econometrica, 1982, 50(4): 987-1007.",
    "[4] Bollerslev T. Generalized Autoregressive Conditional Heteroskedasticity[J]. Journal of Econometrics, 1986, 31(3): 307-327.",
    "[5] J P Morgan, Reuters. RiskMetrics Technical Document[R]. New York: Morgan Guaranty Trust Company, 1996.",
    "[6] Kupiec P H. Techniques for Verifying the Accuracy of Risk Measurement Models[J]. The Journal of Derivatives, 1995, 3(2): 73-84.",
    "[7] Hull J C. Options, Futures, and Other Derivatives[M]. 11th ed. Harlow: Pearson, 2022.",
    "[8] DuckDB Foundation. DuckDB Documentation[EB/OL]. [2026-04-18]. https://duckdb.org/docs/.",
    "[9] Apache Software Foundation. Apache Spark Documentation[EB/OL]. [2026-04-18]. https://spark.apache.org/docs/latest/.",
    "[10] Ramírez S. FastAPI Documentation[EB/OL]. [2026-04-18]. https://fastapi.tiangolo.com/.",
    "[11] Snowflake Inc. Streamlit Documentation[EB/OL]. [2026-04-18]. https://docs.streamlit.io/.",
    "[12] 中国金融期货交易所. 日统计[EB/OL]. [2026-04-18]. http://www.cffex.com.cn/rtj/.",
    "[13] 中国人民银行调查统计司. 金融统计数据报告[EB/OL]. [2026-04-18]. https://www.pbc.gov.cn/diaochatongjisi/116219/116225/index.html.",
    "[14] 国家统计局. 数据发布[EB/OL]. [2026-04-18]. https://www.stats.gov.cn/sj/zxfb/.",
    "[15] Chen X, Hu Y. Volatility forecasts of stock index futures in China and the US: A hybrid LSTM approach[J]. PLoS ONE, 2022, 17(7): e0271595.",
    "[16] Lu F, Ma F, Bouri E, et al. Do commodity futures have a steering effect on the spot stock market in China? New evidence from volatility forecasting[J]. International Review of Financial Analysis, 2024, 94: 103262.",
    "[17] Zeng Q, Zhang J, Zhong J. China's futures market volatility and sectoral stock market volatility prediction[J]. Energy Economics, 2024, 132: 107429.",
    "[18] Mühleisen H, Raasveldt M. DuckDB: An Embeddable Analytical Database[C]. New York: Association for Computing Machinery, 2019.",
    "[19] Mühleisen H, Raasveldt M. Data Management for Data Science: Towards Embedded Analytics[C]. CIDR, 2020.",
    "[20] 何婧萱. 基于机器学习的 A 股市场股指期货价格预测研究[D]. 杭州: 浙江大学, 2024.",
    "[21] 朱峰. 基于深度学习的金融市场高频交易数据波动率预测的研究与应用[D]. 上海: 东华大学, 2023.",
]
ref_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "参考文献")
thanks_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "致谢")
ref_paragraphs = [p for p in doc.paragraphs[ref_idx + 1 : thanks_idx] if p.text.strip()]
for idx, reference in enumerate(refs):
    if idx >= len(ref_paragraphs):
        break
    paragraph = ref_paragraphs[idx]
    clear_runs(paragraph)
    run = paragraph.add_run(reference)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(23)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Cm(0.74)
    pf.first_line_indent = Cm(-0.74)
    set_run_font(run, "宋体", "Times New Roman", 12, False)

# Body citations use square-bracket superscript according to the school rule.
citation_pattern = re.compile(r"\[[0-9]+(?:[,，、-][0-9]+)*\]")
body_start_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "1 绪论")
ref_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "参考文献")
caption_start_pattern = re.compile(r"^[图表]\d+[.．]\d+")
for paragraph in doc.paragraphs[body_start_idx:ref_idx]:
    text = paragraph.text
    if not citation_pattern.search(text):
        continue
    if paragraph.style.name.startswith("Heading") or caption_start_pattern.match(text.strip()):
        continue
    has_code_font = any(
        run.font.name == "Consolas"
        or (
            run._element.rPr is not None
            and run._element.rPr.rFonts is not None
            and run._element.rPr.rFonts.get(qn("w:eastAsia")) == "Consolas"
        )
        for run in paragraph.runs
    )
    if has_code_font:
        continue
    parts = []
    last = 0
    for match in citation_pattern.finditer(text):
        if match.start() > last:
            parts.append((text[last : match.start()], False))
        parts.append((match.group(0), True))
        last = match.end()
    if last < len(text):
        parts.append((text[last:], False))
    clear_runs(paragraph)
    for value, is_citation in parts:
        if not value:
            continue
        run = paragraph.add_run(value)
        set_run_font(run, "宋体", "Times New Roman", 12, False)
        if is_citation:
            run.font.superscript = True

doc.save(OUT)
print(OUT)

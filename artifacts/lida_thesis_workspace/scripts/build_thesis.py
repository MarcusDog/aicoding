from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


WORKSPACE = Path(__file__).resolve().parents[1]
PROJECT_ROOT = WORKSPACE.parents[1]
DELIVERABLES = WORKSPACE / "deliverables"
SOURCE_PATH = WORKSPACE / "drafts" / "thesis_source.md"
WORKING_TEMPLATE = DELIVERABLES / "thesis-working.docx"
OUTPUT_DOCX = DELIVERABLES / "thesis-final.docx"
GENERATED_DIR = WORKSPACE / "evidence" / "generated_assets"


@dataclass
class Section:
    heading: str
    lines: list[str]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the Shanghai Lida thesis DOCX from source markdown.")
    parser.add_argument("--source", default=str(SOURCE_PATH))
    parser.add_argument("--template", default=str(WORKING_TEMPLATE))
    parser.add_argument("--output", default=str(OUTPUT_DOCX))
    parser.add_argument("--skip-postprocess", action="store_true")
    return parser.parse_args()


def parse_source(path: Path) -> tuple[dict[str, str], list[Section]]:
    metadata: dict[str, str] = {}
    sections: list[Section] = []
    lines = path.read_text(encoding="utf-8").splitlines()

    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            break
        if ":" in line and not line.startswith("#"):
            key, value = line.split(":", 1)
            metadata[key.strip()] = value.strip()
            index += 1
            continue
        break

    current_heading: str | None = None
    current_lines: list[str] = []
    for raw_line in lines[index:]:
        if raw_line.startswith("# "):
            if current_heading is not None:
                sections.append(Section(current_heading, current_lines))
            current_heading = raw_line[2:].strip()
            current_lines = []
        else:
            current_lines.append(raw_line)
    if current_heading is not None:
        sections.append(Section(current_heading, current_lines))

    return metadata, sections


def get_font(size: int, bold: bool = False, english: bool = False) -> tuple[str, str, Pt, bool]:
    east_asia = "Times New Roman" if english else "宋体"
    ascii_font = "Times New Roman" if english else "Times New Roman"
    return east_asia, ascii_font, Pt(size), bold


def set_run_font(run, east_asia: str, ascii_font: str, size: Pt, bold: bool = False) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def apply_paragraph_font(paragraph, east_asia: str, ascii_font: str, size: Pt, bold: bool = False) -> None:
    if not paragraph.runs:
        run = paragraph.add_run("")
        set_run_font(run, east_asia, ascii_font, size, bold)
        return
    for run in paragraph.runs:
        set_run_font(run, east_asia, ascii_font, size, bold)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = normal.paragraph_format
    pf.line_spacing = Pt(23)
    pf.first_line_indent = Pt(24)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for style_name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        if style_name not in doc.styles:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = Pt(23)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.first_line_indent = Pt(0)


def replace_paragraph_text(paragraph, text: str, english: bool = False, bold: bool = False) -> None:
    paragraph.text = text
    east_asia, ascii_font, size, default_bold = get_font(12, bold=bold, english=english)
    apply_paragraph_font(paragraph, east_asia, ascii_font, size, bold or default_bold)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def clear_body(doc: Document) -> None:
    start_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name == "Heading 1":
            start_index = index
            break
    if start_index is None:
        return
    for paragraph in list(doc.paragraphs[start_index:]):
        delete_paragraph(paragraph)


def clean_text_for_thesis(text: str) -> str:
    return (
        text.replace("`", "")
        .replace("\\lambda", "λ")
        .replace("\\", "")
    )


def pinyin_name(name: str) -> str:
    mapping = {
        "郭子健": "Guo Zijian",
        "马利平": "Ma Liping",
    }
    return mapping.get(name, name)


def set_cover_fields(doc: Document, meta: dict[str, str], cn_abstract: str, en_abstract: str) -> None:
    doc.paragraphs[1].runs[1].text = meta["STUDENT_ID"]
    cover_title_cn = "基于大数据的中国金融期货交易分析系统\n设计与实现"
    cover_title_en = "Design and Implementation of a Big Data Based\nChinese Financial Futures Trading Analysis System"
    replace_paragraph_text(doc.paragraphs[8], cover_title_cn)
    replace_paragraph_text(doc.paragraphs[12], f"学院名称  ：   {meta['COLLEGE']}     ")
    replace_paragraph_text(doc.paragraphs[13], f"专业名称 ：{meta['MAJOR']}")
    replace_paragraph_text(doc.paragraphs[14], f"学生姓名 ：       {meta['STUDENT_NAME']}       ")
    replace_paragraph_text(doc.paragraphs[15], f"指导教师 ：        {meta['SUPERVISOR']}        ")
    replace_paragraph_text(doc.paragraphs[17], meta["DATE_CN"])

    replace_paragraph_text(doc.paragraphs[22], cover_title_en, english=True)
    replace_paragraph_text(doc.paragraphs[26], f"Candidate：{pinyin_name(meta['STUDENT_NAME'])}", english=True)
    replace_paragraph_text(doc.paragraphs[27], f"Supervisor：{pinyin_name(meta['SUPERVISOR'])}", english=True)
    replace_paragraph_text(doc.paragraphs[32], meta["DATE_EN"], english=True)

    replace_paragraph_text(doc.paragraphs[55], meta["TITLE_CN"])
    replace_paragraph_text(doc.paragraphs[57], "摘要")
    replace_paragraph_text(doc.paragraphs[59], clean_text_for_thesis(cn_abstract))
    replace_paragraph_text(doc.paragraphs[61], clean_text_for_thesis(f"关键词：{meta['KEYWORDS_CN']}"))

    replace_paragraph_text(doc.paragraphs[78], clean_text_for_thesis(en_abstract), english=True)
    replace_paragraph_text(doc.paragraphs[80], clean_text_for_thesis(f"Key words:{meta['KEYWORDS_EN']}"), english=True)

    for idx in [55, 57, 59, 61]:
        apply_paragraph_font(doc.paragraphs[idx], "宋体", "Times New Roman", Pt(12), bold=(idx == 57))
    for idx in [76, 78, 80]:
        apply_paragraph_font(doc.paragraphs[idx], "Times New Roman", "Times New Roman", Pt(12), bold=(idx == 76))


def extract_abstracts(sections: list[Section]) -> tuple[str, str, list[Section]]:
    cn = ""
    en = ""
    rest: list[Section] = []
    for section in sections:
        joined = "\n".join([line for line in section.lines if line.strip()]).strip()
        if section.heading == "中文摘要":
            cn = joined
        elif section.heading == "English Abstract":
            en = joined
        else:
            rest.append(section)
    return cn, en, rest


def body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = Pt(23)
    run = paragraph.add_run(clean_text_for_thesis(text))
    set_run_font(run, "宋体", "Times New Roman", Pt(12), False)


def heading_paragraph(doc: Document, text: str, level: int) -> None:
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(23)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", "Times New Roman", Pt(18 if level == 1 else 14 if level == 2 else 12), True)


def caption_paragraph(doc: Document, text: str, above: bool = False) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(23)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(clean_text_for_thesis(text))
    set_run_font(run, "宋体", "Times New Roman", Pt(12), False)


def reference_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.hanging_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = Pt(23)
    run = paragraph.add_run(clean_text_for_thesis(text))
    set_run_font(run, "宋体", "Times New Roman", Pt(12), False)


def code_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.line_spacing = Pt(14)
    run = paragraph.add_run(text)
    set_run_font(run, "Consolas", "Consolas", Pt(8), False)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = text
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(18)
        apply_paragraph_font(paragraph, "宋体", "Times New Roman", Pt(10), bold)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if edge_data:
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))
        elif element is not None:
            tc_borders.remove(element)


def apply_three_line_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 12, "space": 0, "color": "000000"} if row_index == 0 else None,
                bottom={"val": "single", "sz": 12, "space": 0, "color": "000000"} if row_index in (0, len(table.rows) - 1) else None,
            )


def parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        rows.append(parts)
    if len(rows) >= 2 and all(part.startswith("---") for part in rows[1]):
        rows.pop(1)
    return rows


def try_font(names: list[str], size: int) -> ImageFont.FreeTypeFont:
    for name in names:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_multiline_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill: str) -> None:
    left, top, width, height = box
    lines = text.split("\n")
    line_height = font.size + 8
    start_y = top + (height - line_height * len(lines)) // 2
    for idx, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=font)
        text_width = bbox[2] - bbox[0]
        x = left + (width - text_width) // 2
        y = start_y + idx * line_height
        draw.text((x, y), line, font=font, fill=fill)


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, font) -> None:
    x, y, width, height = xy
    draw.rounded_rectangle((x, y, x + width, y + height), radius=18, fill="#fffaf3", outline="#b8842f", width=3)
    draw_multiline_text(draw, (x, y, width, height), text, font, "#1d2a33")


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line((start, end), fill="#7d3d1f", width=4)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - 16 * direction, y2 - 8), (x2 - 16 * direction, y2 + 8)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 8, y2 - 16 * direction), (x2 + 8, y2 - 16 * direction)]
    draw.polygon(points, fill="#7d3d1f")


def ensure_diagram_pngs() -> dict[str, Path]:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    title_font = try_font(["msyh.ttc", "simhei.ttf"], 34)
    box_font = try_font(["msyh.ttc", "simhei.ttf"], 24)
    mapping: dict[str, Path] = {}

    specs = {
        "fig_4_1_system_architecture": {
            "title": "系统总体架构图",
            "boxes": [
                (60, 120, 200, 90, "CFFEX\n日统计"),
                (60, 300, 200, 90, "PBC / NBS\n宏观数据"),
                (360, 210, 240, 110, "数据采集层\ningest"),
                (700, 210, 240, 110, "数据存储层\nDuckDB / Parquet"),
                (1040, 120, 240, 90, "分析服务层\nFastAPI"),
                (1040, 300, 240, 90, "展示层\nStreamlit"),
            ],
            "arrows": [
                ((260, 165), (360, 250)),
                ((260, 345), (360, 280)),
                ((600, 265), (700, 265)),
                ((940, 250), (1040, 165)),
                ((940, 280), (1040, 345)),
            ],
        },
        "fig_4_2_data_flow": {
            "title": "数据处理流程图",
            "boxes": [
                (50, 210, 200, 80, "原始网页/XML"),
                (300, 210, 220, 80, "标准化清洗"),
                (570, 210, 220, 80, "主力连续构建"),
                (860, 110, 220, 80, "趋势/波动率"),
                (860, 310, 220, 80, "相关性/VaR"),
                (1150, 210, 220, 80, "页面与论文配图"),
            ],
            "arrows": [
                ((250, 250), (300, 250)),
                ((520, 250), (570, 250)),
                ((790, 250), (860, 150)),
                ((790, 250), (860, 350)),
                ((1080, 150), (1150, 250)),
                ((1080, 350), (1150, 250)),
            ],
        },
        "fig_3_1_data_model": {
            "title": "核心数据模型图",
            "boxes": [
                (40, 100, 250, 130, "contracts\ninstrument_id\nproduct_id\ncontract_month"),
                (40, 310, 250, 160, "futures_daily\ntrading_date\nopen/high/low/close\nvolume/open_interest"),
                (380, 310, 260, 160, "main_contract_daily\ndaily_return\nrolling_vol_20\ndrawdown"),
                (760, 310, 250, 160, "macro_series\nseries_name\nobservation_date\nvalue"),
                (1130, 310, 250, 160, "analysis_snapshot\nVaR\ntrend_signal\nrisk_level"),
            ],
            "arrows": [
                ((165, 230), (165, 310)),
                ((290, 390), (380, 390)),
                ((640, 390), (760, 390)),
                ((1010, 390), (1130, 390)),
            ],
        },
    }

    for stem, spec in specs.items():
        output = GENERATED_DIR / f"{stem}.png"
        if not output.exists():
            image = Image.new("RGB", (1450, 550), "#f8f4ed")
            draw = ImageDraw.Draw(image)
            draw.text((580, 30), spec["title"], font=title_font, fill="#1d2a33")
            for box in spec["boxes"]:
                draw_box(draw, box[:4], box[4], box_font)
            for start, end in spec["arrows"]:
                draw_arrow(draw, start, end)
            image.save(output)
        mapping[stem] = output
    return mapping


def resolve_figure_path(name: str, generated_diagrams: dict[str, Path]) -> Path:
    stem = name.strip()
    if stem in generated_diagrams:
        return generated_diagrams[stem]

    candidate_dirs = [
        PROJECT_ROOT / "artifacts" / "thesis" / "screenshots",
        PROJECT_ROOT / "artifacts" / "thesis" / "figures",
        PROJECT_ROOT / "artifacts" / "thesis" / "diagrams",
    ]
    for directory in candidate_dirs:
        for ext in [".png", ".jpg", ".jpeg"]:
            candidate = directory / f"{stem}{ext}"
            if candidate.exists():
                return candidate
    raise FileNotFoundError(f"Could not resolve figure asset: {name}")


def add_figure(doc: Document, asset_name: str, caption: str, generated_diagrams: dict[str, Path]) -> None:
    path = resolve_figure_path(asset_name, generated_diagrams)
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(5.6))
    caption_paragraph(doc, caption)


def add_code(doc: Document, relative_path: str, start: int, end: int, caption: str) -> None:
    caption_paragraph(doc, caption)
    target = PROJECT_ROOT / relative_path
    lines = target.read_text(encoding="utf-8").splitlines()
    for number, line in enumerate(lines[start - 1 : end], start=start):
        code_paragraph(doc, f"{number:>3}: {line.rstrip()}")


def add_equation(doc: Document, formula: str, number: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(23)
    run = paragraph.add_run(f"OMATH::{clean_text_for_thesis(formula)}")
    set_run_font(run, "Cambria Math", "Cambria Math", Pt(12), False)

    number_paragraph = doc.add_paragraph(style="Normal")
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_paragraph.paragraph_format.first_line_indent = Pt(0)
    number_paragraph.paragraph_format.line_spacing = Pt(23)
    run = number_paragraph.add_run(clean_text_for_thesis(number))
    set_run_font(run, "宋体", "Times New Roman", Pt(12), False)


def add_table(doc: Document, caption: str, rows: list[list[str]]) -> None:
    caption_paragraph(doc, caption)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            set_cell_text(table.cell(row_index, col_index), value, bold=(row_index == 0))
    apply_three_line_table(table)


def render_sections(doc: Document, sections: list[Section], generated_diagrams: dict[str, Path]) -> None:
    first_chapter = True
    current_top_heading = ""

    for section in sections:
        if not first_chapter:
            doc.add_page_break()
        first_chapter = False
        current_top_heading = section.heading
        heading_paragraph(doc, section.heading, 1)

        index = 0
        while index < len(section.lines):
            line = section.lines[index].rstrip()
            stripped = line.strip()
            if not stripped:
                index += 1
                continue
            if stripped.startswith("## "):
                heading_paragraph(doc, stripped[3:].strip(), 2)
                index += 1
                continue
            if stripped.startswith("### "):
                heading_paragraph(doc, stripped[4:].strip(), 3)
                index += 1
                continue
            if stripped.startswith("[[FIGURE:"):
                payload = stripped[len("[[FIGURE:") : -2]
                asset_name, caption = payload.split("|", 1)
                add_figure(doc, asset_name, caption, generated_diagrams)
                index += 1
                continue
            if stripped.startswith("[[CODE:"):
                payload = stripped[len("[[CODE:") : -2]
                relative_path, start, end, caption = payload.split("|", 3)
                add_code(doc, relative_path, int(start), int(end), caption)
                index += 1
                continue
            if stripped.startswith("[[EQUATION:"):
                payload = stripped[len("[[EQUATION:") : -2]
                formula, number = payload.split("|", 1)
                add_equation(doc, formula, number)
                index += 1
                continue
            if stripped.startswith("[[TABLE:"):
                caption = stripped[len("[[TABLE:") : -2]
                table_lines: list[str] = []
                index += 1
                while index < len(section.lines) and section.lines[index].strip().startswith("|"):
                    table_lines.append(section.lines[index].strip())
                    index += 1
                add_table(doc, caption, parse_markdown_table(table_lines))
                continue
            if current_top_heading == "参考文献":
                reference_paragraph(doc, stripped)
            else:
                body_paragraph(doc, stripped)
            index += 1


def run_postprocess(output_path: Path) -> None:
    subprocess.run(
        [sys.executable, str(WORKSPACE / "scripts" / "postprocess_thesis.py"), "--docx", str(output_path)],
        check=True,
        cwd=WORKSPACE,
    )


def main() -> None:
    args = parse_args()
    source = Path(args.source).resolve()
    template = Path(args.template).resolve()
    output = Path(args.output).resolve()
    output.parent.mkdir(parents=True, exist_ok=True)

    metadata, sections = parse_source(source)
    cn_abstract, en_abstract, body_sections = extract_abstracts(sections)
    generated_diagrams = ensure_diagram_pngs()

    shutil.copy2(template, output)
    doc = Document(str(output))
    configure_styles(doc)
    set_cover_fields(doc, metadata, cn_abstract, en_abstract)
    clear_body(doc)
    render_sections(doc, body_sections, generated_diagrams)
    doc.save(output)

    if not args.skip_postprocess:
        run_postprocess(output)

    print(f"Built thesis DOCX: {output}")


if __name__ == "__main__":
    main()

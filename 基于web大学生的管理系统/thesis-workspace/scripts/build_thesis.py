from __future__ import annotations

import re
from pathlib import Path

import win32com.client
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from thesis_content import (
    ACKNOWLEDGEMENTS,
    CN_ABSTRACT,
    CN_KEYWORDS,
    EN_ABSTRACT,
    EN_KEYWORDS,
    MANUSCRIPT,
    METADATA,
    REFERENCES,
    WORKSPACE,
)


BASE_DOC = WORKSPACE / "deliverables" / "thesis-working.docx"
OUTPUT_DOC = WORKSPACE / "deliverables" / "王冠超-基于Web的大学生志愿者管理系统设计与开发-论文终稿.docx"
OUTPUT_PDF = WORKSPACE / "deliverables" / "王冠超-基于Web的大学生志愿者管理系统设计与开发-论文终稿.pdf"
AUDIT_FILE = WORKSPACE / "audit" / "thesis-build-audit.txt"
SCRIPT_DIR = Path(__file__).resolve().parent
NSMAP = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def set_run_font(run, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size: float = 12, bold: bool = False):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def replace_text_keep_paragraph(paragraph, text: str, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size: float = 12):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size=size)


def replace_cover_field(paragraph, label: str, value: str, underline_spaces: int = 12):
    clear_paragraph(paragraph)
    label_run = paragraph.add_run(f"{label}  ：")
    set_run_font(label_run, east_asia="宋体", ascii_font="Times New Roman", size=12)
    value_run = paragraph.add_run(f"{value}{' ' * underline_spaces}")
    value_run.font.underline = True
    set_run_font(value_run, east_asia="宋体", ascii_font="Times New Roman", size=12)


def replace_multiline_title(paragraph, lines: list[str], east_asia: str, size: float, bold: bool):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(28)
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, east_asia=east_asia, ascii_font="Times New Roman", size=size, bold=bold)
        if idx < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def delete_block_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def paragraph_has_section_break(paragraph: Paragraph) -> bool:
    return paragraph._element.find(".//w:sectPr", NSMAP) is not None


def ensure_page_break_before(paragraph: Paragraph):
    paragraph.paragraph_format.page_break_before = True


def format_body_paragraph(paragraph, first_line_indent: bool = True):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(24) if first_line_indent else Pt(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(23)
    for run in paragraph.runs:
        set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12)


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12)
    format_body_paragraph(p, first_line_indent=True)
    return p


def insert_page_break_before(anchor: Paragraph):
    p = anchor.insert_paragraph_before()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p


def insert_body_paragraph_before(anchor: Paragraph, text: str):
    p = anchor.insert_paragraph_before()
    run = p.add_run(text)
    set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12)
    format_body_paragraph(p, first_line_indent=True)
    return p


def insert_center_title_before(anchor: Paragraph, text: str, size: float = 16, bold: bool = True, east_asia: str = "黑体", ascii_font: str = "Times New Roman"):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(23)
    run = p.add_run(text)
    set_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)
    return p


def insert_keywords_before(anchor: Paragraph, text: str, english: bool = False):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if english else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(0) if english else Pt(24)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(23)
    run = p.add_run(text)
    set_run_font(
        run,
        east_asia="Times New Roman" if english else "宋体",
        ascii_font="Times New Roman",
        size=12,
    )
    return p


def add_center_title(doc: Document, text: str, size: float = 16, bold: bool = True, east_asia: str = "黑体", ascii_font: str = "Times New Roman"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(23)
    run = p.add_run(text)
    set_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(23)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=16, bold=True)
    elif level == 2:
        set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=14, bold=True)
    else:
        set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=12, bold=True)
    return p


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("left", "right", "insideV", "insideH"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")
    for edge in ("top", "bottom"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:color"), "000000")


def set_cell_bottom_border(cell, width: str = "8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), width)
    bottom.set(qn("w:color"), "000000")


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]):
    add_caption(doc, caption, caption_type="table")
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for i, head in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = head
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(20)
        for run in p.runs:
            set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=10.5, bold=True)
        set_cell_bottom_border(cell)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(20)
            for run in p.runs:
                set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=10.5)
            if r_idx == len(rows):
                set_cell_bottom_border(cell)
    doc.add_paragraph()


def add_caption(doc: Document, text: str, caption_type: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(text)
    set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=10.5)
    if caption_type == "table":
        return p
    return p


def add_figure(doc: Document, path: Path, caption: str, width_cm: float):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption, caption_type="figure")


def add_code_block(doc: Document, caption: str, code: str):
    add_caption(doc, caption, caption_type="table")
    for line in code.splitlines():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(14)
        run = p.add_run(line.rstrip())
        set_run_font(run, east_asia="Consolas", ascii_font="Consolas", size=9)
    doc.add_paragraph()


def add_equation_placeholder(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(f"[[EQ]]{text}")
    set_run_font(run, east_asia="Cambria Math", ascii_font="Cambria Math", size=11)
    return p


def normalize_equation_text(text: str) -> str:
    mappings = {
        r"\eta_j = \frac{N_{approved,j}}{Q_j} \times 100\%": "η_j = (N_(approved,j))/(Q_j) × 100%",
        r"H_{total} = \sum_{i=1}^{n} h_i \cdot \delta_i": "H_(total) = Σ_(i=1)^n h_i · δ_i",
    }
    return mappings.get(text, text)


def set_borderless_table(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def add_equation_block(doc: Document, text: str, number_label: str):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_borderless_table(table)

    left_p = table.cell(0, 0).paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    mid_p = table.cell(0, 1).paragraphs[0]
    mid_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mid_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    mid_p.paragraph_format.line_spacing = Pt(20)
    eq_run = mid_p.add_run(f"[[EQ]]{normalize_equation_text(text)}")
    set_run_font(eq_run, east_asia="Cambria Math", ascii_font="Cambria Math", size=11)

    right_p = table.cell(0, 2).paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    right_p.paragraph_format.line_spacing = Pt(20)
    num_run = right_p.add_run(number_label)
    set_run_font(num_run, east_asia="宋体", ascii_font="Times New Roman", size=10.5)
    doc.add_paragraph()


def rewrite_cover(doc: Document):
    replacements = {
        1: f"学号：{METADATA['student_no']}",
        12: f"学院名称  ：   {METADATA['college']}",
        13: f"专业名称  ：{METADATA['major']}",
        14: f"学生姓名  ：       {METADATA['student_name']}",
        15: f"指导教师  ：        {METADATA['supervisor']}",
        17: METADATA["date_cn"],
        26: f"Candidate：{METADATA['candidate_en']}",
        27: f"Supervisor：{METADATA['supervisor_en']}",
        32: METADATA["date_en"],
    }
    for idx, text in replacements.items():
        replace_text_keep_paragraph(doc.paragraphs[idx], text)
    replace_multiline_title(doc.paragraphs[8], ["基于Web的大学生志愿者管理系统", "设计与开发"], east_asia="黑体", size=18, bold=False)
    replace_cover_field(doc.paragraphs[12], "学院名称", METADATA["college"])
    replace_cover_field(doc.paragraphs[13], "专业名称", METADATA["major"])
    replace_cover_field(doc.paragraphs[14], "学生姓名", METADATA["student_name"])
    replace_cover_field(doc.paragraphs[15], "指导教师", METADATA["supervisor"])
    replace_multiline_title(
        doc.paragraphs[22],
        ["Design and Development of a Web-Based", "College Student Volunteer Management System"],
        east_asia="Times New Roman",
        size=16,
        bold=True,
    )


def repair_statement_table(doc: Document):
    table = doc.tables[0]
    table.cell(0, 2).text = METADATA["major"]
    table.cell(1, 2).text = METADATA["student_no"]
    table.cell(2, 2).text = ""
    table.cell(3, 2).text = "年    月    日"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph.paragraph_format.line_spacing = Pt(20)
                for run in paragraph.runs:
                    set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=10.5)


def prune_old_content(doc: Document):
    body = doc._element.body
    children = list(body.iterchildren())
    section_breaks = [
        child for child in children
        if child.tag == qn("w:p") and child.find(".//w:sectPr", NSMAP) is not None
    ]
    if len(section_breaks) < 3:
        raise ValueError("Template body does not contain three section boundaries")

    section2_break = section_breaks[1]
    section3_break = section_breaks[2]

    remove_between = False
    for child in list(body.iterchildren()):
        if child is section2_break:
            remove_between = True
            continue
        if child is section3_break:
            remove_between = False
            continue
        if remove_between:
            delete_block_element(child)

    remove_body = False
    for child in list(body.iterchildren()):
        if child is section3_break:
            remove_body = True
            continue
        if child.tag == qn("w:sectPr"):
            continue
        if remove_body:
            delete_block_element(child)


def locate_front_anchor(doc: Document) -> Paragraph:
    section_breaks = [paragraph for paragraph in doc.paragraphs if paragraph_has_section_break(paragraph)]
    if len(section_breaks) < 3:
        raise ValueError("Template is missing required section breaks")
    return section_breaks[2]


def build_front_matter(anchor: Paragraph):
    insert_center_title_before(anchor, METADATA["title_cn"], size=16, bold=True, east_asia="黑体")
    insert_center_title_before(anchor, "摘  要", size=15, bold=True, east_asia="黑体")
    insert_body_paragraph_before(anchor, CN_ABSTRACT)
    insert_keywords_before(anchor, f"关键词：{CN_KEYWORDS}")
    insert_page_break_before(anchor)

    insert_center_title_before(anchor, "BACHELOR'S DEGREE THESIS OF SHANGHAI LIDA UNIVERSITY", size=14, bold=False, east_asia="Times New Roman", ascii_font="Times New Roman")
    insert_center_title_before(anchor, "ABSTRACT", size=15, bold=True, east_asia="Times New Roman", ascii_font="Times New Roman")
    title_p = anchor.insert_paragraph_before()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(0)
    title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    title_p.paragraph_format.line_spacing = Pt(20)
    title_run = title_p.add_run(METADATA["title_en"])
    set_run_font(title_run, east_asia="Times New Roman", ascii_font="Times New Roman", size=12)
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(23)
    run = p.add_run(EN_ABSTRACT)
    set_run_font(run, east_asia="Times New Roman", ascii_font="Times New Roman", size=12)
    insert_keywords_before(anchor, f"Key words: {EN_KEYWORDS}", english=True)
    insert_page_break_before(anchor)

    insert_center_title_before(anchor, "目  录", size=15, bold=True, east_asia="黑体")
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("[[TOC]]")
    set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12)


def build_body(doc: Document):
    first_chapter = True
    for chapter in MANUSCRIPT:
        if not first_chapter:
            doc.add_page_break()
        first_chapter = False
        add_heading(doc, chapter["title"], 1)
        chapter_no_match = re.match(r"(\d+)", chapter["title"])
        chapter_no = chapter_no_match.group(1) if chapter_no_match else "0"
        equation_index = 0
        for element in chapter["elements"]:
            kind = element["type"]
            if kind == "heading":
                level = element["level"]
                mapped = 2 if level == 2 else 3
                add_heading(doc, element["text"], mapped)
            elif kind == "paragraph":
                add_body_paragraph(doc, element["text"])
            elif kind == "figure":
                add_figure(doc, element["path"], element["caption"], element["width_cm"])
            elif kind == "table":
                add_table(doc, element["caption"], element["headers"], element["rows"])
            elif kind == "code":
                add_code_block(doc, element["caption"], element["code"])
            elif kind == "equation":
                equation_index += 1
                add_equation_block(doc, element["text"], f"（{chapter_no}-{equation_index}）")
            else:
                raise ValueError(f"Unknown element type: {kind}")


def build_references(doc: Document):
    doc.add_page_break()
    add_heading(doc, "参考文献", 1)
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.first_line_indent = Pt(-24)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.left_indent = Pt(24)
        run = p.add_run(ref)
        set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=10.5)


def build_acknowledgements(doc: Document):
    doc.add_page_break()
    add_heading(doc, "致  谢", 1)
    for paragraph in ACKNOWLEDGEMENTS:
        add_body_paragraph(doc, paragraph)


def postprocess_with_word(path: Path, pdf_path: Path):
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(path))
        wd_header_footer_primary = 1
        wd_page_number_style_arabic = 0
        wd_page_number_style_uppercase_roman = 1
        wd_align_page_number_center = 1
        toc_range = doc.Range(0, doc.Content.End)
        toc_find = toc_range.Find
        toc_find.ClearFormatting()
        toc_find.Text = "[[TOC]]"
        if toc_find.Execute():
            toc_range.Text = ""
            doc.TablesOfContents.Add(
                Range=toc_range,
                UseHeadingStyles=True,
                UpperHeadingLevel=1,
                LowerHeadingLevel=3,
                UseFields=True,
                TableID="",
                RightAlignPageNumbers=True,
                IncludePageNumbers=True,
                AddedStyles="",
                UseHyperlinks=True,
                HidePageNumbersInWeb=False,
                UseOutlineLevels=True,
            )

        search_start = 0
        while True:
            eq_range = doc.Range(search_start, doc.Content.End)
            eq_find = eq_range.Find
            eq_find.ClearFormatting()
            eq_find.Text = "[[EQ]]"
            if not eq_find.Execute():
                break
            para_range = eq_range.Paragraphs(1).Range
            raw = para_range.Text.replace("\r", "").replace("\x07", "").strip()
            expr = raw.replace("[[EQ]]", "", 1)
            para_range.Text = expr
            para_range.OMaths.Add(para_range)
            para_range.OMaths(1).BuildUp()
            search_start = para_range.End + 1

        for section_index in (1, 2):
            if section_index > doc.Sections.Count:
                break
            section = doc.Sections(section_index)
            section.Headers(wd_header_footer_primary).LinkToPrevious = False
            section.Footers(wd_header_footer_primary).LinkToPrevious = False
            section.Headers(wd_header_footer_primary).Range.Text = ""
            section.Footers(wd_header_footer_primary).Range.Text = ""

        if doc.Sections.Count >= 3:
            front_section = doc.Sections(3)
            front_section.Headers(wd_header_footer_primary).LinkToPrevious = False
            front_section.Footers(wd_header_footer_primary).LinkToPrevious = False
            front_section.Headers(wd_header_footer_primary).Range.Text = ""
            front_footer = front_section.Footers(wd_header_footer_primary)
            front_footer.Range.Text = ""
            while front_footer.PageNumbers.Count > 0:
                front_footer.PageNumbers(1).Delete()
            front_footer.PageNumbers.RestartNumberingAtSection = True
            front_footer.PageNumbers.NumberStyle = wd_page_number_style_uppercase_roman
            front_footer.PageNumbers.StartingNumber = 1
            front_footer.PageNumbers.Add(wd_align_page_number_center, True)

        if doc.Sections.Count >= 4:
            body_section = doc.Sections(4)
            body_section.Headers(wd_header_footer_primary).LinkToPrevious = False
            body_section.Footers(wd_header_footer_primary).LinkToPrevious = False
            body_footer = body_section.Footers(wd_header_footer_primary)
            body_footer.Range.Text = ""
            while body_footer.PageNumbers.Count > 0:
                body_footer.PageNumbers(1).Delete()
            body_footer.PageNumbers.RestartNumberingAtSection = True
            body_footer.PageNumbers.NumberStyle = wd_page_number_style_arabic
            body_footer.PageNumbers.StartingNumber = 1
            body_footer.PageNumbers.Add(wd_align_page_number_center, True)

        if doc.TablesOfContents.Count > 0:
            doc.TablesOfContents(1).Update()
        doc.Fields.Update()
        doc.Save()
        doc.ExportAsFixedFormat(str(pdf_path), 17)
        doc.Close()
    finally:
        word.Quit()


def build_audit(doc: Document):
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    chars = len(re.sub(r"\s+", "", text))
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    AUDIT_FILE.write_text(
        "\n".join(
            [
                f"Output: {OUTPUT_DOC}",
                f"PDF: {OUTPUT_PDF}",
                f"Paragraph count: {len(doc.paragraphs)}",
                f"Table count: {len(doc.tables)}",
                f"Non-space chars: {chars}",
                f"Chinese chars: {chinese_chars}",
            ]
        ),
        encoding="utf-8",
    )


def main():
    assets_script = SCRIPT_DIR / "generate_thesis_assets.py"
    if assets_script.exists():
        namespace: dict[str, object] = {"__file__": str(assets_script)}
        exec(assets_script.read_text(encoding="utf-8"), namespace)
        namespace["generate_assets"]()

    doc = Document(str(BASE_DOC))
    rewrite_cover(doc)
    repair_statement_table(doc)
    ensure_page_break_before(doc.paragraphs[37])
    prune_old_content(doc)
    build_front_matter(locate_front_anchor(doc))
    build_body(doc)
    build_references(doc)
    build_acknowledgements(doc)
    doc.save(str(OUTPUT_DOC))
    postprocess_with_word(OUTPUT_DOC, OUTPUT_PDF)
    final_doc = Document(str(OUTPUT_DOC))
    build_audit(final_doc)


if __name__ == "__main__":
    main()

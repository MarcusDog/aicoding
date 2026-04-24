from __future__ import annotations

import json
import re
import shutil

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from thesis_content import ACKNOWLEDGMENT, CODE_SNIPPETS, DOCUMENT_ITEMS, FIGURES, METADATA, PROJECT_ROOT, REFERENCES, WORKSPACE_ROOT


DELIVERABLES = WORKSPACE_ROOT / "deliverables"
OUTPUT_DOCX = DELIVERABLES / "financial-fraud-thesis-draft.docx"
OUTPUT_MD = WORKSPACE_ROOT / "drafts" / "thesis_content.md"
AUDIT_MD = WORKSPACE_ROOT / "audit" / "content-audit.md"


def set_run_font(run, east_asia: str, ascii_font: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def replace_paragraph_text(paragraph, text: str, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size_pt: float = 12.0, bold: bool = False, align: int | None = None) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size_pt=size_pt, bold=bold)
    if align is not None:
        paragraph.alignment = align


def remove_body_from_first_heading(doc: Document) -> None:
    started = False
    body = doc.element.body
    for child in list(body.iterchildren()):
        if child.tag == qn("w:p"):
            texts = child.xpath(".//w:t/text()")
            joined = "".join(texts).strip()
            if joined == "1 绪　论":
                started = True
        if started and child.tag != qn("w:sectPr"):
            body.remove(child)


def format_body_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = Pt(23)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def add_citation_aware_runs(paragraph, text: str, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size_pt: float = 12.0) -> None:
    citation_pattern = re.compile(r"(\[[0-9,\-]+\])")
    parts = citation_pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        set_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size_pt=size_pt, bold=False)
        if citation_pattern.fullmatch(part):
            run.font.superscript = True


def add_body_text(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    format_body_paragraph(paragraph)
    add_citation_aware_runs(paragraph, text)


def add_chapter(doc: Document, title: str, page_break: bool = True) -> None:
    if page_break:
        doc.add_page_break()
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(title)
    set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size_pt=18, bold=True)


def add_section(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = Pt(23)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size_pt=14, bold=True)


def add_table_title(doc: Document, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = Pt(23)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(caption)
    set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size_pt=12, bold=True)


def add_figure_caption(doc: Document, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = Pt(23)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(caption)
    set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size_pt=12, bold=True)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def make_three_line_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size_pt=12, bold=True)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size_pt=12, bold=False)
    border_top = {"val": "single", "sz": 12, "color": "000000", "space": 0}
    border_mid = {"val": "single", "sz": 8, "color": "000000", "space": 0}
    border_none = {"val": "nil"}
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            kwargs = {"left": border_none, "right": border_none, "insideV": border_none}
            if row_idx == 0:
                kwargs["top"] = border_top
                kwargs["bottom"] = border_mid
            elif row_idx == len(table.rows) - 1:
                kwargs["bottom"] = border_top
            set_cell_border(cell, **kwargs)


def add_figure(doc: Document, figure_key: str) -> None:
    figure = FIGURES[figure_key]
    figure_path = figure["path"]
    target_dir = WORKSPACE_ROOT / ("screenshots" if "screenshots" in figure_path.parts else "charts")
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / figure_path.name
    if not target_path.exists():
        shutil.copy2(figure_path, target_path)
    doc.add_page_break()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(target_path), width=Cm(13.2))
    add_figure_caption(doc, figure["caption"])


def add_code_block(doc: Document, code_key: str) -> None:
    item = CODE_SNIPPETS[code_key]
    add_table_title(doc, item["title"])
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = item["text"]
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    cell._tc.get_or_add_tcPr().append(shading)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing = Pt(16)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        for run in paragraph.runs:
            set_run_font(run, east_asia="等线", ascii_font="Consolas", size_pt=10.5, bold=False)


def add_formula(doc: Document, equation: str, number: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.text = f"{{EQ:{equation}}}"
    right.text = number
    for paragraph in left.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, east_asia="Times New Roman", ascii_font="Times New Roman", size_pt=12)
    for paragraph in right.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size_pt=12)
    border_none = {"val": "nil"}
    set_cell_border(left, top=border_none, bottom=border_none, left=border_none, right=border_none)
    set_cell_border(right, top=border_none, bottom=border_none, left=border_none, right=border_none)


def load_table_data(key: str) -> tuple[list[str], list[list[str]]]:
    if key == "requirements":
        return (
            ["需求类别", "具体要求"],
            [
                ["离线训练", "支持多模型实验、阈值搜索、指标导出和图表生成"],
                ["在线推理", "支持单笔预测、批量预测、模式查询和健康检查"],
                ["准实时监控", "支持Kafka消息回放、Spark微批评分与结果落盘"],
                ["可视化展示", "支持总览、预测、监控和论文配图查看"],
                ["论文支撑", "输出截图、图表、模型文件和实验统计结果"],
            ],
        )
    if key == "dataset_overview":
        metadata = json.loads((PROJECT_ROOT / "data" / "raw" / "creditcard.metadata.json").read_text(encoding="utf-8"))
        return (
            ["项目", "数值"],
            [
                ["样本总数", str(metadata["rows"])],
                ["特征列数", str(metadata["columns"])],
                ["目标字段", metadata["target_column"]],
                ["欺诈样本数", str(metadata["fraud_count"])],
                ["欺诈样本占比", f"{metadata['fraud_ratio']:.4%}"],
                ["数据来源", "OpenML creditcard"],
            ],
        )
    if key == "split_summary":
        rows = []
        for name in ["train", "valid", "test"]:
            frame = pd.read_csv(PROJECT_ROOT / "data" / "processed" / "splits" / f"{name}.csv")
            fraud = int(frame["Class"].sum())
            rows.append([name, str(len(frame)), str(fraud), f"{fraud / len(frame):.4%}"])
        return (["数据集", "样本数", "欺诈样本数", "欺诈占比"], rows)
    if key == "api_endpoints":
        return (
            ["接口", "方法", "功能说明"],
            [
                ["/health", "GET", "返回服务状态与模型加载状态"],
                ["/metrics", "GET", "返回最佳实验指标"],
                ["/schema", "GET", "返回模型所需特征模式"],
                ["/predict", "POST", "执行单笔交易欺诈预测"],
                ["/predict/batch", "POST", "执行批量交易欺诈预测"],
            ],
        )
    if key == "model_results":
        frame = pd.read_csv(PROJECT_ROOT / "artifacts" / "metrics" / "experiment_results.csv")
        frame = frame[["name", "pr_auc", "precision", "recall", "f1", "fpr"]]
        rows = []
        for _, row in frame.iterrows():
            rows.append([str(row["name"]), f"{row['pr_auc']:.4f}", f"{row['precision']:.4f}", f"{row['recall']:.4f}", f"{row['f1']:.4f}", f"{row['fpr']:.4f}"])
        return (["实验方案", "PR-AUC", "Precision", "Recall", "F1", "FPR"], rows)
    if key == "feature_importance_top":
        frame = pd.read_csv(PROJECT_ROOT / "artifacts" / "metrics" / "feature_importance.csv").head(10)
        rows = [[str(feature), str(importance)] for feature, importance in frame.values.tolist()]
        return (["特征名", "重要性"], rows)
    if key == "functional_tests":
        return (
            ["测试项", "测试方式", "预期结果", "测试结论"],
            [
                ["健康检查接口", "GET /health", "返回status=ok", "通过"],
                ["指标查询接口", "GET /metrics", "返回PR-AUC等指标", "通过"],
                ["特征模式接口", "GET /schema", "返回29个特征及示例值", "通过"],
                ["单笔预测功能", "离线样例页面", "输出风险分、阈值和解释原因", "通过"],
                ["批量预测功能", "上传CSV文件", "生成批量预测结果表", "通过"],
                ["准实时监控功能", "Kafka+Spark回放", "生成最新批次汇总和监控图", "通过"],
                ["自动化测试", "pytest", "接口与展示逻辑断言成功", "通过"],
            ],
        )
    raise KeyError(key)


def write_markdown_source() -> None:
    lines = [
        f"# {METADATA['title_cn']}",
        "",
        "## 中文摘要",
        METADATA["abstract_cn"],
        "",
        f"关键词：{METADATA['keywords_cn']}",
        "",
        "## English Abstract",
        METADATA["abstract_en"],
        "",
        f"Keywords: {METADATA['keywords_en']}",
        "",
    ]
    for item in DOCUMENT_ITEMS:
        if item["type"] == "chapter":
            lines.extend([f"# {item['title']}", ""])
        elif item["type"] == "section":
            lines.extend([f"## {item['title']}", ""])
        elif item["type"] == "paragraph":
            lines.extend([item["text"], ""])
        elif item["type"] == "figure":
            figure = FIGURES[item["key"]]
            lines.extend([f"![{figure['caption']}]({figure['path']})", "", figure["caption"], ""])
        elif item["type"] == "table":
            lines.extend([f"{{TABLE:{item['key']}|{item['caption']}}}", ""])
        elif item["type"] == "code":
            snippet = CODE_SNIPPETS[item["key"]]
            lines.extend([snippet["title"], f"```{snippet['language']}", snippet["text"], "```", ""])
        elif item["type"] == "formula":
            lines.extend([f"{{FORMULA:{item['equation']}|{item['number']}}}", ""])
    lines.extend(["# 参考文献", ""])
    lines.extend(REFERENCES)
    lines.extend(["", "# 致谢", "", ACKNOWLEDGMENT, ""])
    OUTPUT_MD.write_text("\n".join(lines), encoding="utf-8")


def replace_front_matter(doc: Document) -> None:
    mapping = {
        2: f"学号：{METADATA['student_id']}",
        9: METADATA["title_cn"],
        13: f"学院名称  ：   {METADATA['college_cn']}",
        14: f"专业名称\u00a0：{METADATA['major_cn']}",
        15: f"学生姓名\u00a0：       {METADATA['student_name_cn']}",
        16: f"指导教师\u00a0：        {METADATA['supervisor_cn']}",
        18: METADATA["date_cn"],
        23: METADATA["title_en"],
        27: f"Candidate：{METADATA['student_name_en']}",
        28: f"Supervisor：{METADATA['supervisor_en']}",
        33: METADATA["date_en"],
        56: METADATA["title_cn"],
        60: METADATA["abstract_cn"],
        62: f"关键词：{METADATA['keywords_cn']}",
        79: METADATA["abstract_en"],
        81: f"Key words: {METADATA['keywords_en']}",
    }
    for index, text in mapping.items():
        paragraph = doc.paragraphs[index - 1]
        if index in {9, 23, 56}:
            replace_paragraph_text(paragraph, text, east_asia="黑体", ascii_font="Times New Roman", size_pt=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif index in {60, 62}:
            replace_paragraph_text(paragraph, text, east_asia="宋体", ascii_font="Times New Roman", size_pt=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
            paragraph.paragraph_format.line_spacing = Pt(23)
        elif index in {79, 81}:
            replace_paragraph_text(paragraph, text, east_asia="Times New Roman", ascii_font="Times New Roman", size_pt=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            paragraph.paragraph_format.line_spacing = Pt(23)
        else:
            replace_paragraph_text(paragraph, text, east_asia="宋体", ascii_font="Times New Roman", size_pt=12, bold=False)
    if doc.tables:
        declaration_table = doc.tables[0]
        declaration_table.cell(0, 2).text = METADATA["major_cn"]
        declaration_table.cell(1, 2).text = METADATA["student_id"]
        declaration_table.cell(2, 2).text = ""
        for row in declaration_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.line_spacing = Pt(23)
                    for run in paragraph.runs:
                        set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size_pt=12, bold=False)


def add_references(doc: Document) -> None:
    add_chapter(doc, "参考文献")
    for ref in REFERENCES:
        paragraph = doc.add_paragraph()
        format_body_paragraph(paragraph)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run(ref)
        set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size_pt=12)


def add_acknowledgment(doc: Document) -> None:
    add_chapter(doc, "致 谢")
    add_body_text(doc, ACKNOWLEDGMENT)


def build_body(doc: Document) -> None:
    first_chapter = True
    for item in DOCUMENT_ITEMS:
        if item["type"] == "chapter":
            add_chapter(doc, item["title"], page_break=not first_chapter)
            first_chapter = False
        elif item["type"] == "section":
            add_section(doc, item["title"])
        elif item["type"] == "paragraph":
            add_body_text(doc, item["text"])
        elif item["type"] == "table":
            add_table_title(doc, item["caption"])
            headers, rows = load_table_data(item["key"])
            make_three_line_table(doc, headers, rows)
        elif item["type"] == "figure":
            add_figure(doc, item["key"])
        elif item["type"] == "code":
            add_code_block(doc, item["key"])
        elif item["type"] == "formula":
            add_formula(doc, item["equation"], item["number"])


def compute_body_character_count(doc: Document) -> int:
    text = []
    started = False
    for paragraph in doc.paragraphs:
        raw = paragraph.text.strip()
        if raw == "1 绪　论":
            started = True
        if started and raw:
            text.append(raw)
    merged = "".join(text)
    merged = merged.replace("参考文献", "").replace("致 谢", "")
    return len(re.sub(r"\s+", "", merged))


def write_audit(doc: Document) -> None:
    count = compute_body_character_count(doc)
    lines = [
        "# Thesis Content Audit",
        "",
        f"- 正文字数（按字符粗略统计）：{count}",
        f"- 参考文献条目数：{len(REFERENCES)}",
        f"- 插图数量：{sum(1 for item in DOCUMENT_ITEMS if item['type'] == 'figure')}",
        f"- 表格数量：{sum(1 for item in DOCUMENT_ITEMS if item['type'] == 'table')}",
        f"- 代码节选数量：{sum(1 for item in DOCUMENT_ITEMS if item['type'] == 'code')}",
        f"- 公式数量：{sum(1 for item in DOCUMENT_ITEMS if item['type'] == 'formula')}",
        "",
        "## 说明",
        "- 目录、页码与公式专业排版需在Word后处理阶段刷新。",
        "- 指导教师姓名未从现有开题报告中获取，当前封面字段保持空白。",
    ]
    AUDIT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    template_path = WORKSPACE_ROOT / "deliverables" / "thesis-working.docx"
    doc = Document(str(template_path))
    remove_body_from_first_heading(doc)
    replace_front_matter(doc)
    build_body(doc)
    add_references(doc)
    add_acknowledgment(doc)
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))
    write_markdown_source()
    write_audit(Document(str(OUTPUT_DOCX)))
    print(f"Draft DOCX: {OUTPUT_DOCX}")
    print(f"Markdown source: {OUTPUT_MD}")
    print(f"Audit report: {AUDIT_MD}")


if __name__ == "__main__":
    main()

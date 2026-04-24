from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from latex2mathml.converter import convert as latex_to_mathml
from mathml2omml import convert as mathml_to_omml


WORKSPACE = Path(__file__).resolve().parent
PROJECT_ROOT = WORKSPACE.parents[1]
CONTENT_PATH = WORKSPACE / "drafts" / "thesis_content.md"
WORKING_DOCX = WORKSPACE / "deliverables" / "thesis-working.docx"
OUTPUT_DOCX = WORKSPACE / "deliverables" / "潘哲-毕业论文-终稿.docx"


FIGURE_ROOT = PROJECT_ROOT / "evidence" / "figures"
SCREENSHOT_ROOT = PROJECT_ROOT / "evidence" / "screenshots"
BODY_LINE_SPACING = Pt(23)
CHAPTER_SPACE_BEFORE = Pt(18.4)  # 0.8 line when body line spacing is 23 pt.
HALF_LINE_SPACE = Pt(11.5)


TOKEN_RE = re.compile(r"^\[\[(FIGURE|TABLE|FORMULA|CODE):([^|\]]+)\|(.+)\]\]$")
TABLE_DEF_RE = re.compile(r"^\[\[TABLEDEF:([^\]]+)\]\]$")
CODE_DEF_RE = re.compile(r"^\[\[CODEDEF:([^\]]+)\]\]$")
INLINE_MATH_RE = re.compile(r"\$(.+?)\$")

STATIC_TOC_PAGES = {
    "1 绪论": 8,
    "1.1 研究背景": 8,
    "1.2 国内外研究现状": 8,
    "1.3 研究内容与技术路线": 9,
    "1.4 论文结构": 10,
    "2 相关技术与理论基础": 12,
    "2.1 Python 数据处理技术栈": 12,
    "2.2 地图展示与交互技术": 12,
    "2.3 空间数据与外部服务支撑": 13,
    "2.4 评分方法的理论基础": 13,
    "2.5 本章小结": 14,
    "3 数据采集与处理": 15,
    "3.1 数据采集需求分析": 15,
    "3.2 数据来源与采集过程": 15,
    "3.3 数据清洗与标准化处理": 16,
    "3.4 数据质量与需求满足情况": 17,
    "3.5 本章小结": 17,
    "4 杨浦区地摊经济监测与点位决策网站设计": 18,
    "4.1 总体架构设计": 18,
    "4.2 数据流与模块划分": 19,
    "4.3 关键数据结构设计": 19,
    "4.4 页面与交互设计": 20,
    "4.5 本章小结": 21,
    "5 推荐机制与关键功能实现": 22,
    "5.1 特征构建实现": 22,
    "5.2 规则评分与解释机制": 23,
    "5.3 监督学习扩展接口": 24,
    "5.4 页面实现与可视化输出": 25,
    "5.5 本章小结": 26,
    "6 系统测试": 27,
    "6.1 测试目标与测试方法": 27,
    "6.2 数据处理测试": 27,
    "6.3 功能测试与运行验证": 29,
    "6.4 推荐输出与可靠性测试": 30,
    "6.5 本章小结": 32,
    "7 结论与展望": 33,
    "7.1 研究结论": 33,
    "7.2 后续展望": 33,
    "参考文献": 34,
    "致谢": 36,
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the Shanghai Lida thesis DOCX from structured content.")
    parser.add_argument("--content", type=Path, default=CONTENT_PATH)
    parser.add_argument("--template", type=Path, default=WORKING_DOCX)
    parser.add_argument("--output", type=Path, default=OUTPUT_DOCX)
    return parser.parse_args()


def set_run_font(run, east_asia: str, ascii_font: str | None = None, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = ascii_font or east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_runs(paragraph, east_asia: str, ascii_font: str | None = None, size: float | None = None, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)


def normalize_figure_table_refs(text: str) -> str:
    return re.sub(r"([图表])(\d+)\.(\d+)", r"\1\2-\3", text)


def normalize_reference_punctuation(text: str) -> str:
    punctuation_map = str.maketrans(
        {
            "[": "［",
            "]": "］",
            ",": "，",
            ".": "．",
            ":": "：",
            ";": "；",
            "(": "（",
            ")": "）",
            "-": "－",
        }
    )
    parts: list[str] = []
    last = 0
    for match in re.finditer(r"https?://[^\s,，]+", text):
        parts.append(text[last : match.start()].translate(punctuation_map))
        parts.append(match.group(0))
        last = match.end()
    parts.append(text[last:].translate(punctuation_map))
    return "".join(parts)


def bookmark_name(title: str, index: int) -> str:
    slug = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "_", title).strip("_")
    return f"bm_{index}_{slug[:32]}"


def build_heading_bookmarks(content_lines: list[str]) -> dict[str, str]:
    titles: list[str] = []
    for line in content_lines:
        stripped = line.strip()
        if stripped.startswith("## "):
            titles.append(stripped[3:].strip())
        elif stripped.startswith("### "):
            titles.append(stripped[4:].strip())
    titles.extend(["参考文献", "致谢"])
    return {title: bookmark_name(title, index) for index, title in enumerate(titles, start=1)}


def add_bookmark(paragraph, name: str) -> None:
    match = re.match(r"bm_(\d+)_", name)
    bookmark_id = match.group(1) if match else "0"
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def append_text_run(paragraph, text: str, east_asia: str, ascii_font: str | None = None, size: float | None = None, bold: bool | None = None) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)


def latex_to_omml_element(latex: str):
    normalized = latex.replace("\\\\", "\\")
    mathml = latex_to_mathml(normalized)
    omml = mathml_to_omml(mathml)
    root = parse_xml(f'<root xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{omml}</root>')
    return root[0]


def add_text_with_inline_math(paragraph, text: str, east_asia: str, ascii_font: str | None = None, size: float | None = None, bold: bool | None = None) -> None:
    last = 0
    for match in INLINE_MATH_RE.finditer(text):
        append_text_run(paragraph, text[last : match.start()], east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)
        paragraph._p.append(latex_to_omml_element(match.group(1)))
        last = match.end()
    append_text_run(paragraph, text[last:], east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)


def add_internal_hyperlink(paragraph, text: str, anchor: str, east_asia: str, size: float, bold: bool) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    def styled_run(run_text: str | None = None, tab: bool = False) -> OxmlElement:
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Times New Roman")
        fonts.set(qn("w:hAnsi"), "Times New Roman")
        fonts.set(qn("w:eastAsia"), east_asia)
        rpr.append(fonts)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        rpr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "none")
        rpr.append(underline)
        size_node = OxmlElement("w:sz")
        size_node.set(qn("w:val"), str(int(size * 2)))
        rpr.append(size_node)
        size_cs = OxmlElement("w:szCs")
        size_cs.set(qn("w:val"), str(int(size * 2)))
        rpr.append(size_cs)
        if bold:
            bold_node = OxmlElement("w:b")
            rpr.append(bold_node)
        run.append(rpr)
        if tab:
            run.append(OxmlElement("w:tab"))
        else:
            text_node = OxmlElement("w:t")
            if run_text and (run_text.startswith(" ") or run_text.endswith(" ")):
                text_node.set(qn("xml:space"), "preserve")
            text_node.text = run_text or ""
            run.append(text_node)
        return run

    title, page = text.split("\t", 1)
    hyperlink.append(styled_run(title))
    hyperlink.append(styled_run(tab=True))
    hyperlink.append(styled_run(page))
    paragraph._p.append(hyperlink)


def delete_body_from(doc: Document, start_index: int) -> None:
    body = doc._body._element
    start_element = doc.paragraphs[start_index]._element
    remove = False
    for child in list(body):
        if child == start_element:
            remove = True
        if remove and child.tag != qn("w:sectPr"):
            body.remove(child)


def replace_toc(doc: Document, content_lines: list[str], bookmark_map: dict[str, str]) -> None:
    toc_index = next((idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "目　录"), None)
    if toc_index is None:
        return

    body = doc._body._element
    toc_paragraph = doc.paragraphs[toc_index]
    toc_paragraph.text = "目　录"
    toc_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_paragraph.paragraph_format.first_line_indent = Cm(0)
    toc_paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
    toc_paragraph.paragraph_format.space_before = Pt(0)
    toc_paragraph.paragraph_format.space_after = Pt(0)
    set_paragraph_runs(toc_paragraph, east_asia="黑体", ascii_font="Times New Roman", size=18, bold=True)
    for paragraph in list(doc.paragraphs[toc_index + 1 :]):
        if paragraph._element in body:
            body.remove(paragraph._element)

    entries: list[tuple[int, str]] = []
    for line in content_lines:
        stripped = line.strip()
        if stripped.startswith("## "):
            entries.append((1, stripped[3:].strip()))
        elif stripped.startswith("### "):
            entries.append((2, stripped[4:].strip()))
    entries.extend([(1, "参考文献"), (1, "致谢")])

    anchor = toc_paragraph
    for level, title in entries:
        page = STATIC_TOC_PAGES.get(title, "")
        paragraph = doc.add_paragraph()
        paragraph._p.getparent().remove(paragraph._p)
        anchor._p.addnext(paragraph._p)
        anchor = paragraph

        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0 if level == 1 else 0.74)
        paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(14.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        add_internal_hyperlink(
            paragraph,
            f"{title}\t{page}",
            bookmark_map[title],
            east_asia="黑体" if level == 1 else "宋体",
            size=14 if level == 1 else 12,
            bold=level == 1,
        )


def ensure_rfonts(style, east_asia: str, ascii_font: str) -> None:
    style.font.name = ascii_font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.color.rgb = RGBColor(0, 0, 0)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    ensure_rfonts(normal, "宋体", "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = BODY_LINE_SPACING
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    heading_specs = {
        "Heading 1": (18, WD_ALIGN_PARAGRAPH.CENTER, CHAPTER_SPACE_BEFORE, HALF_LINE_SPACE),
        "Heading 2": (14, WD_ALIGN_PARAGRAPH.LEFT, HALF_LINE_SPACE, HALF_LINE_SPACE),
        "Heading 3": (12, WD_ALIGN_PARAGRAPH.LEFT, Pt(0), Pt(0)),
    }
    for style_name, (size, alignment, space_before, space_after) in heading_specs.items():
        style = doc.styles[style_name]
        ensure_rfonts(style, "黑体", "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = space_before
        style.paragraph_format.space_after = space_after
        style.paragraph_format.line_spacing = BODY_LINE_SPACING


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)


def read_sections(path: Path) -> dict[str, list[str]]:
    text = path.read_text(encoding="utf-8").splitlines()
    sections: dict[str, list[str]] = {}
    current = None
    for line in text:
        if line.startswith("# "):
            current = line[2:].strip()
            sections[current] = []
            continue
        if current is not None:
            sections[current].append(line)
    return sections


def parse_metadata(lines: list[str]) -> dict[str, str]:
    meta: dict[str, str] = {}
    for line in lines:
        if "：" in line:
            key, value = line.split("：", 1)
            meta[key.strip()] = value.strip()
    return meta


def parse_defs(lines: list[str], pattern: re.Pattern[str]) -> dict[str, list[list[str]]]:
    defs: dict[str, list[list[str]]] = {}
    current: str | None = None
    buffer: list[list[str]] = []
    for line in lines:
        match = pattern.match(line)
        if match:
            if current is not None:
                defs[current] = buffer
            current = match.group(1)
            buffer = []
            continue
        if current is not None:
            if line.strip():
                buffer.append([part.strip() for part in line.split("|")])
    if current is not None:
        defs[current] = buffer
    return defs


def parse_code_defs(lines: list[str]) -> dict[str, list[str]]:
    defs: dict[str, list[str]] = {}
    current: str | None = None
    buffer: list[str] = []
    for line in lines:
        match = CODE_DEF_RE.match(line)
        if match:
            if current is not None:
                defs[current] = buffer
            current = match.group(1)
            buffer = []
            continue
        if current is not None:
            buffer.append(line.rstrip("\n"))
    if current is not None:
        defs[current] = buffer
    return defs


def split_references(lines: list[str]) -> list[str]:
    return [line.strip() for line in lines if line.strip()]


def set_cover(doc: Document, meta: dict[str, str]) -> None:
    replacements = {
        1: f"学号：{meta['学号']}",
        8: meta["题目（中文）"],
        12: f"学院名称  ：   {meta['学院']}",
        13: f"专业名称 ：{meta['专业']}",
        14: f"学生姓名 ：       {meta['学生姓名']}",
        15: f"指导教师 ：        {meta['指导教师']}",
        17: meta["日期"].replace("年", "  年  ").replace("月", "  月  ").replace("日", "  日"),
        22: meta["题目（英文）"],
        26: f"Candidate: {meta['学生姓名（拼音）']}",
        27: f"Supervisor: {meta['指导教师（拼音）']}",
        32: meta["英文日期"],
        55: meta["题目（中文）"],
    }
    for index, text in replacements.items():
        paragraph = doc.paragraphs[index]
        if paragraph.runs:
            paragraph.runs[0].text = text
            for extra in paragraph.runs[1:]:
                extra.text = ""
        else:
            paragraph.text = text

    student_id = doc.paragraphs[1]
    for child_name in ("w:ind", "w:tabs"):
        child = student_id._p.pPr.find(qn(child_name))
        if child is not None:
            student_id._p.pPr.remove(child)
    student_id.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    student_id.paragraph_format.first_line_indent = Cm(0)
    student_id.paragraph_format.left_indent = Cm(0)
    student_id.paragraph_format.right_indent = Cm(0)
    set_paragraph_runs(student_id, east_asia="黑体", ascii_font="Times New Roman", size=10.5, bold=False)


def set_abstracts(doc: Document, sections: dict[str, list[str]]) -> None:
    cn_lines = [line for line in sections["中文摘要"] if line.strip()]
    en_lines = [line for line in sections["English Abstract"] if line.strip()]

    cn_abstract = "\n".join(cn_lines[:-1])
    cn_keywords = cn_lines[-1].replace("关键词：", "关键词：")
    en_abstract = "\n".join(en_lines[:-1])
    en_keywords = en_lines[-1]

    for index, text in [(59, cn_abstract), (61, cn_keywords), (78, en_abstract), (80, en_keywords)]:
        paragraph = doc.paragraphs[index]
        paragraph.text = text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if index in (59, 78) else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0.74 if index in (59, 78) else 0)
        paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
        if index in (59, 61):
            set_paragraph_runs(paragraph, east_asia="宋体", ascii_font="Times New Roman", size=12)
        else:
            set_paragraph_runs(paragraph, east_asia="宋体", ascii_font="Times New Roman", size=12)


def add_caption(doc: Document, text: str, above: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(normalize_figure_table_refs(text))
    set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=12, bold=True)


def clear_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["left", "right", "insideV"]:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")
    for edge in ["top", "bottom", "insideH"]:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_table(doc: Document, caption: str, rows: list[list[str]]) -> None:
    add_caption(doc, caption, above=True)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    clear_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = Pt(15.75)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=10.5, bold=False)
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str) -> None:
    figure_path = FIGURE_ROOT / filename
    if not figure_path.exists():
        figure_path = SCREENSHOT_ROOT / filename
    if not figure_path.exists():
        raise FileNotFoundError(f"Figure not found: {filename}")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(figure_path), width=Cm(14.5))
    add_caption(doc, caption)


def add_formula(doc: Document, number: str, expr: str) -> None:
    eq = doc.add_paragraph()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.first_line_indent = Cm(0)
    eq.paragraph_format.line_spacing = 1.0
    eq.paragraph_format.space_before = Pt(0)
    eq.paragraph_format.space_after = Pt(2)
    eq._p.append(latex_to_omml_element(expr))

    no = doc.add_paragraph()
    no.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    no.paragraph_format.first_line_indent = Cm(0)
    no.paragraph_format.line_spacing = BODY_LINE_SPACING
    run = no.add_run(f"（{number}）")
    set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12)


def add_code_block(doc: Document, caption: str, lines: list[str]) -> None:
    add_caption(doc, caption, above=True)
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = Pt(18)
        run = paragraph.add_run(line)
        set_run_font(run, east_asia="等线", ascii_font="Consolas", size=10.5)
    doc.add_paragraph()


def add_body(
    doc: Document,
    content_lines: list[str],
    table_defs: dict[str, list[list[str]]],
    code_defs: dict[str, list[str]],
    bookmark_map: dict[str, str],
) -> None:
    for line in content_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped == "[[PAGEBREAK]]":
            doc.add_page_break()
            continue
        token = TOKEN_RE.match(stripped)
        if token:
            token_type, token_id, payload = token.groups()
            if token_type == "FIGURE":
                add_figure(doc, token_id, payload)
            elif token_type == "TABLE":
                add_table(doc, payload, table_defs[token_id])
            elif token_type == "FORMULA":
                add_formula(doc, token_id, payload)
            elif token_type == "CODE":
                add_code_block(doc, payload, code_defs[token_id])
            continue

        if stripped.startswith("## "):
            doc.add_page_break()
            paragraph = doc.add_paragraph(stripped[3:].strip(), style="Heading 1")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.left_indent = Cm(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = CHAPTER_SPACE_BEFORE
            paragraph.paragraph_format.space_after = HALF_LINE_SPACE
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            set_paragraph_runs(paragraph, east_asia="黑体", ascii_font="Times New Roman", size=18, bold=True)
            add_bookmark(paragraph, bookmark_map[stripped[3:].strip()])
            continue
        if stripped.startswith("### "):
            title = stripped[4:].strip()
            style = "Heading 2" if re.match(r"^\d+\.\d+\s", title) else "Heading 3"
            size = 14 if style == "Heading 2" else 12
            paragraph = doc.add_paragraph(title, style=style)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.left_indent = Cm(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if style == "Heading 2":
                paragraph.paragraph_format.space_before = HALF_LINE_SPACE
                paragraph.paragraph_format.space_after = HALF_LINE_SPACE
            else:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            set_paragraph_runs(paragraph, east_asia="黑体", ascii_font="Times New Roman", size=size, bold=True)
            add_bookmark(paragraph, bookmark_map[title])
            continue

        paragraph = doc.add_paragraph()
        paragraph.style = doc.styles["Normal"]
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_text_with_inline_math(
            paragraph,
            normalize_figure_table_refs(stripped),
            east_asia="宋体",
            ascii_font="Times New Roman",
            size=12,
        )


def add_references(doc: Document, refs: list[str], bookmark_map: dict[str, str]) -> None:
    doc.add_page_break()
    heading = doc.add_paragraph("参考文献", style="Heading 1")
    heading.paragraph_format.first_line_indent = Cm(0)
    heading.paragraph_format.left_indent = Cm(0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = CHAPTER_SPACE_BEFORE
    heading.paragraph_format.space_after = HALF_LINE_SPACE
    heading.paragraph_format.line_spacing = BODY_LINE_SPACING
    set_paragraph_runs(heading, east_asia="黑体", ascii_font="Times New Roman", size=18, bold=True)
    add_bookmark(heading, bookmark_map["参考文献"])
    for ref in refs:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(-0.74)
        paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(normalize_reference_punctuation(ref))
        set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12)


def add_acknowledgement(doc: Document, lines: list[str], bookmark_map: dict[str, str]) -> None:
    doc.add_page_break()
    heading = doc.add_paragraph("致谢", style="Heading 1")
    heading.paragraph_format.first_line_indent = Cm(0)
    heading.paragraph_format.left_indent = Cm(0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = CHAPTER_SPACE_BEFORE
    heading.paragraph_format.space_after = HALF_LINE_SPACE
    heading.paragraph_format.line_spacing = BODY_LINE_SPACING
    set_paragraph_runs(heading, east_asia="黑体", ascii_font="Times New Roman", size=18, bold=True)
    add_bookmark(heading, bookmark_map["致谢"])
    for line in lines:
        if not line.strip():
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(line.strip())
        set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12)


def main() -> None:
    args = parse_args()
    sections = read_sections(args.content)
    metadata = parse_metadata(sections["元数据"])
    table_defs = parse_defs(sections["表格定义"], TABLE_DEF_RE)
    code_defs = parse_code_defs(sections["代码片段定义"])
    refs = split_references(sections["参考文献"])
    body_lines = sections["正文"]
    ack_lines = sections["致谢"]
    bookmark_map = build_heading_bookmarks(body_lines)

    doc = Document(args.template)
    configure_styles(doc)
    configure_sections(doc)
    set_cover(doc, metadata)
    set_abstracts(doc, sections)
    delete_body_from(doc, 131)
    replace_toc(doc, body_lines, bookmark_map)
    add_body(doc, body_lines, table_defs, code_defs, bookmark_map)
    add_references(doc, refs, bookmark_map)
    add_acknowledgement(doc, ack_lines, bookmark_map)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()

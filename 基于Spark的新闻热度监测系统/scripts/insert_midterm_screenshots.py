from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = next(path for path in (ROOT / "output" / "doc").glob("*.docx") if "_before_screenshots" not in path.stem)
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_before_screenshots.docx")

SCREENSHOTS = [
    ("dashboard-page.png", "图1 新闻监测可视化大屏页面"),
    ("news-page.png", "图2 新闻查询与导出页面"),
    ("event-page.png", "图3 事件聚类与关联分析页面"),
]


def set_run_font(run, size: float = 12.0) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def format_caption(paragraph, text: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, 12.0)


def format_image_paragraph(paragraph, image_path: Path) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.7))


def find_target_paragraph(doc: Document):
    table = doc.tables[0]
    cell = table.cell(4, 4)
    target = None
    texts = [p.text.strip() for p in cell.paragraphs]
    if any("图1 新闻监测可视化大屏页面" in text for text in texts):
        return cell, None
    for paragraph in cell.paragraphs:
        if "1.7" in paragraph.text and "测试与阶段性成果整理" in paragraph.text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError("未找到 1.7 小节起始段落，无法插入截图。")
    return cell, target


def insert_before(target, paragraph_type: str, payload):
    new_paragraph = target.insert_paragraph_before("")
    if paragraph_type == "image":
        format_image_paragraph(new_paragraph, payload)
    elif paragraph_type == "caption":
        format_caption(new_paragraph, payload)
    else:
        raise ValueError(paragraph_type)
    return new_paragraph


def main() -> None:
    if BACKUP_PATH.exists():
        BACKUP_PATH.unlink()

    doc = Document(str(DOC_PATH))
    cell, target = find_target_paragraph(doc)
    if target is None:
        raise RuntimeError("截图似乎已经插入，且当前脚本要求从无截图版本重新执行。")

    DOC_PATH.replace(BACKUP_PATH)
    doc = Document(str(BACKUP_PATH))
    cell, target = find_target_paragraph(doc)

    items: list[tuple[str, object]] = []
    for filename, caption in SCREENSHOTS:
        image_path = ROOT / "output" / "doc" / filename
        if not image_path.exists():
            raise FileNotFoundError(f"missing screenshot: {image_path}")
        items.append(("image", image_path))
        items.append(("caption", caption))

    current_target = target
    for kind, payload in reversed(items):
        current_target = insert_before(current_target, kind, payload)

    doc.save(str(DOC_PATH))
    print(DOC_PATH)
    print(BACKUP_PATH)


if __name__ == "__main__":
    main()

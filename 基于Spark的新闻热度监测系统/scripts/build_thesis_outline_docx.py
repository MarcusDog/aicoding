from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_LEADER, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "output/doc/templates/毕业设计说明书内容格式（空白模板）.docx"
SOURCE = ROOT / "docs/thesis-outline-v1.md"
OUTPUT = ROOT / "output/doc/基于Spark的新闻热度监测系统_论文骨架初稿.docx"
FIGURE_DIR = ROOT / "output/doc/figures"
COVER_TITLE = "基于Spark的新闻热度监测系统"
COVER_DATE = "2025年6月"
HEADER_TEXT = "中北大学2025届毕业设计说明书"
TOC_DATA = ROOT / "output/doc/toc_entries.txt"
ZH_ABSTRACT = (
    "针对新闻信息来源多样、更新频繁和重复转载影响热点判断等问题，本文设计并实现了一个基于 Spark 的新闻热度监测系统。"
    "系统以本地数据集、RSS 源和网页解析结果为数据入口，对新闻标题、正文、来源、发布时间等字段进行标准化处理，"
    "并通过 HTML 清洗、时间归一化、强去重和弱去重提升数据质量。在分析层，系统采用 Spark 批处理思想完成热点主题统计、"
    "关键词趋势分析、基础情感统计、事件聚类和预警记录生成；在服务层，使用 Flask 封装查询、概览、事件详情和数据导出接口；"
    "在展示层，使用 Vue3 与 ECharts 实现可视化大屏、新闻查询、事件分析和服务中心页面。测试结果表明，系统能够完成新闻采集、"
    "清洗、分析、展示和导出的完整流程，具有结构清晰、结果可复核和扩展性较好的特点，可为新闻热点监测和舆情初步分析提供参考。"
)
ZH_KEYWORDS = "Spark，新闻热度，数据清洗，事件聚类，可视化"
EN_TITLE = "Design and Implementation of a News Popularity Monitoring System Based on Spark"
EN_ABSTRACT = (
    "To address the problems of heterogeneous news sources, frequent updates and repeated reposts affecting popularity judgment, "
    "this thesis designs and implements a news popularity monitoring system based on Spark. The system uses local datasets, RSS feeds "
    "and web page parsing results as data inputs. It standardizes fields such as title, content, source and publication time, and improves "
    "data quality through HTML cleaning, time normalization, exact deduplication and near-duplicate detection. In the analysis layer, the "
    "system follows a Spark batch-processing approach to generate hot topic statistics, keyword trends, basic sentiment summaries, event "
    "clusters and alert records. In the service layer, Flask is used to provide APIs for overview data, news search, event details and data "
    "export. In the presentation layer, Vue3 and ECharts are used to build the dashboard, news search page, event analysis page and service "
    "center. The test results show that the system can complete the full process of news collection, cleaning, analysis, visualization and "
    "export, with clear structure, reproducible results and good extensibility."
)
EN_KEYWORDS = "Spark, news popularity, data cleaning, event clustering, visualization"


def clear_document_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def clear_story_part(story_part) -> None:
    element = story_part._element
    for child in list(element):
        element.remove(child)
    story_part.add_paragraph()


def set_east_asia_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name: str, size_pt: int, bold: bool = False) -> None:
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def ensure_spacing_element(container) -> OxmlElement:
    p_pr = container.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        container.append(p_pr)
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    return spacing


def set_style_spacing_lines(style, *, before_lines: int | None = None, after_lines: int | None = None, line_multiple: float | None = None) -> None:
    spacing = ensure_spacing_element(style._element)
    if before_lines is not None:
        spacing.set(qn("w:beforeLines"), str(before_lines))
        spacing.attrib.pop(qn("w:before"), None)
    if after_lines is not None:
        spacing.set(qn("w:afterLines"), str(after_lines))
        spacing.attrib.pop(qn("w:after"), None)
    if line_multiple is not None:
        spacing.set(qn("w:line"), str(int(240 * line_multiple)))
        spacing.set(qn("w:lineRule"), "auto")


def set_paragraph_spacing_lines(paragraph, *, before_lines: int | None = None, after_lines: int | None = None, line_multiple: float | None = None) -> None:
    spacing = ensure_spacing_element(paragraph._element)
    if before_lines is not None:
        spacing.set(qn("w:beforeLines"), str(before_lines))
        spacing.attrib.pop(qn("w:before"), None)
    if after_lines is not None:
        spacing.set(qn("w:afterLines"), str(after_lines))
        spacing.attrib.pop(qn("w:after"), None)
    if line_multiple is not None:
        spacing.set(qn("w:line"), str(int(240 * line_multiple)))
        spacing.set(qn("w:lineRule"), "auto")


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    set_style_font(normal, "宋体", 12, False)
    normal.paragraph_format.first_line_indent = Pt(24)
    set_style_spacing_lines(normal, before_lines=0, after_lines=0, line_multiple=1.5)

    body = ensure_paragraph_style(document, "Thesis Body", "Normal")
    set_style_font(body, "宋体", 12, False)
    body.paragraph_format.first_line_indent = Pt(24)
    set_style_spacing_lines(body, before_lines=0, after_lines=0, line_multiple=1.5)

    heading1 = ensure_paragraph_style(document, "Heading 1", "Normal")
    set_style_font(heading1, "黑体", 15, True)
    heading1.paragraph_format.first_line_indent = Pt(0)
    set_style_spacing_lines(heading1, before_lines=50, after_lines=50, line_multiple=1.5)

    heading2 = ensure_paragraph_style(document, "Heading 2", "Normal")
    set_style_font(heading2, "黑体", 12, True)
    heading2.paragraph_format.first_line_indent = Pt(0)
    set_style_spacing_lines(heading2, before_lines=0, after_lines=0, line_multiple=1.5)

    heading3 = ensure_paragraph_style(document, "Heading 3", "Normal")
    set_style_font(heading3, "黑体", 12, False)
    heading3.paragraph_format.first_line_indent = Pt(0)
    set_style_spacing_lines(heading3, before_lines=0, after_lines=0, line_multiple=1.5)

    ref_style = ensure_paragraph_style(document, "Thesis Reference", "Normal")
    set_style_font(ref_style, "宋体", 12, False)
    ref_style.paragraph_format.first_line_indent = Pt(0)
    set_style_spacing_lines(ref_style, before_lines=0, after_lines=0, line_multiple=1.5)

    caption_style = ensure_paragraph_style(document, "Thesis Caption", "Normal")
    set_style_font(caption_style, "宋体", 10, False)
    caption_style.paragraph_format.first_line_indent = Pt(0)
    set_style_spacing_lines(caption_style, before_lines=0, after_lines=0, line_multiple=1.0)

    cover_main = ensure_paragraph_style(document, "Thesis Cover Main", "Normal")
    set_style_font(cover_main, "方正小标宋简体", 52, False)
    cover_main.paragraph_format.first_line_indent = Pt(0)
    set_style_spacing_lines(cover_main, before_lines=1600, after_lines=0, line_multiple=1.0)

    cover_title = ensure_paragraph_style(document, "Thesis Cover Title", "Normal")
    set_style_font(cover_title, "黑体", 26, True)
    cover_title.paragraph_format.first_line_indent = Pt(0)
    set_style_spacing_lines(cover_title, before_lines=300, after_lines=0, line_multiple=1.0)

    cover_info = ensure_paragraph_style(document, "Thesis Cover Info", "Normal")
    set_style_font(cover_info, "宋体", 16, False)
    cover_info.paragraph_format.first_line_indent = Pt(0)
    cover_info.paragraph_format.left_indent = Pt(150)
    set_style_spacing_lines(cover_info, before_lines=100, after_lines=0, line_multiple=1.0)

    cover_date = ensure_paragraph_style(document, "Thesis Cover Date", "Normal")
    set_style_font(cover_date, "宋体", 18, False)
    cover_date.paragraph_format.first_line_indent = Pt(0)
    set_style_spacing_lines(cover_date, before_lines=240, after_lines=0, line_multiple=1.0)

    abstract_title = ensure_paragraph_style(document, "Thesis Abstract Title", "Normal")
    set_style_font(abstract_title, "黑体", 15, True)
    abstract_title.paragraph_format.first_line_indent = Pt(0)
    set_style_spacing_lines(abstract_title, before_lines=120, after_lines=50, line_multiple=1.5)

    abstract_heading = ensure_paragraph_style(document, "Thesis Abstract Heading", "Normal")
    set_style_font(abstract_heading, "黑体", 12, True)
    abstract_heading.paragraph_format.first_line_indent = Pt(0)
    set_style_spacing_lines(abstract_heading, before_lines=50, after_lines=50, line_multiple=1.5)

    abstract_body = ensure_paragraph_style(document, "Thesis Abstract Body", "Normal")
    set_style_font(abstract_body, "宋体", 12, False)
    abstract_body.paragraph_format.first_line_indent = Pt(24)
    set_style_spacing_lines(abstract_body, before_lines=0, after_lines=0, line_multiple=1.5)

    abstract_keywords = ensure_paragraph_style(document, "Thesis Abstract Keywords", "Normal")
    set_style_font(abstract_keywords, "宋体", 12, False)
    abstract_keywords.paragraph_format.first_line_indent = Pt(0)
    set_style_spacing_lines(abstract_keywords, before_lines=50, after_lines=0, line_multiple=1.5)

    toc_title = ensure_paragraph_style(document, "Thesis TOC Title", "Normal")
    set_style_font(toc_title, "黑体", 14, True)
    toc_title.paragraph_format.first_line_indent = Pt(0)
    set_style_spacing_lines(toc_title, before_lines=240, after_lines=180, line_multiple=1.0)

    for level, indent in ((1, 0), (2, 18), (3, 36)):
        toc_style = ensure_paragraph_style(document, f"Thesis TOC {level}", "Normal")
        set_style_font(toc_style, "宋体", 12, False)
        toc_style.paragraph_format.first_line_indent = Pt(0)
        toc_style.paragraph_format.left_indent = Pt(indent)
        toc_style.paragraph_format.right_indent = Pt(0)
        set_style_spacing_lines(toc_style, before_lines=0, after_lines=0, line_multiple=1.5)


def enable_update_fields(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def ensure_paragraph_style(document: Document, style_name: str, base_style_name: str):
    try:
        return document.styles[style_name]
    except KeyError:
        style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles[base_style_name]
        return style


def add_heading(document: Document, text: str, level: int, first_h1: bool) -> None:
    if level == 1 and not first_h1:
        document.add_page_break()

    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if text in {"参考文献", "致谢"} else WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing_lines(paragraph, before_lines=50 if level == 1 else 0, after_lines=50 if level == 1 else 0, line_multiple=1.5)
    paragraph.add_run(text)


def set_paragraph_bottom_border(paragraph, color: str = "000000", size: str = "8") -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def set_section_page_number(section, fmt: str, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)
    pg_num_type.set(qn("w:start"), str(start))


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
    set_east_asia_font(run, "宋体", 10, False)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在打开文档或导出 PDF 时自动更新"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)
    set_east_asia_font(run, "宋体", 12, False)


def configure_section_layout(section) -> None:
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(2.5)
    section.footer_distance = Cm(1.8)


def load_toc_entries() -> list[tuple[int, str, str]]:
    if not TOC_DATA.exists():
        return []
    entries: list[tuple[int, str, str]] = []
    for line in TOC_DATA.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        level, title, page = line.split("|", 2)
        entries.append((int(level), title, page))
    return entries


def add_toc_entry(document: Document, level: int, title: str, page: str) -> None:
    paragraph = document.add_paragraph(style=f"Thesis TOC {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.tab_stops.add_tab_stop(Pt(396), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    set_paragraph_spacing_lines(paragraph, before_lines=0, after_lines=0, line_multiple=1.5)
    run = paragraph.add_run(f"{title}\t{page}")
    set_east_asia_font(run, "宋体", 12, False)


def configure_running_section(section, number_fmt: str, start: int) -> None:
    configure_section_layout(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_story_part(section.header)
    clear_story_part(section.footer)
    header_paragraph = section.header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.paragraph_format.line_spacing = 1.0
    header_run = header_paragraph.add_run(HEADER_TEXT)
    set_east_asia_font(header_run, "黑体", 12, False)
    set_paragraph_bottom_border(header_paragraph)
    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.line_spacing = 1.0
    add_page_number_field(footer_paragraph)
    set_section_page_number(section, number_fmt, start)


def configure_front_matter_section(section) -> None:
    configure_section_layout(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_story_part(section.header)
    clear_story_part(section.footer)
    header_paragraph = section.header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.paragraph_format.line_spacing = 1.0
    header_run = header_paragraph.add_run(HEADER_TEXT)
    set_east_asia_font(header_run, "黑体", 12, False)
    set_paragraph_bottom_border(header_paragraph)


def configure_cover_section(section) -> None:
    configure_section_layout(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_story_part(section.header)
    clear_story_part(section.footer)


def add_cover(document: Document) -> None:
    top = document.add_paragraph(style="Thesis Cover Main")
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = top.add_run("毕业设计说明书")
    set_east_asia_font(title_run, "方正小标宋简体", 52, False)

    thesis_title = document.add_paragraph(style="Thesis Cover Title")
    thesis_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    thesis_run = thesis_title.add_run(COVER_TITLE)
    set_east_asia_font(thesis_run, "黑体", 26, True)

    info_lines = [
        "班    级：__________            学号：__________",
        "姓    名：__________",
        "学    院：软件学院",
        "专    业：软件工程",
        "指导教师：__________",
    ]
    for idx, line in enumerate(info_lines):
        paragraph = document.add_paragraph(style="Thesis Cover Info")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing_lines(paragraph, before_lines=100 if idx == 0 else 60, after_lines=0, line_multiple=1.0)
        run = paragraph.add_run(line)
        set_east_asia_font(run, "宋体", 16, False)

    date_paragraph = document.add_paragraph(style="Thesis Cover Date")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(COVER_DATE)
    set_east_asia_font(date_run, "宋体", 18, False)


def add_abstract_pages(document: Document) -> None:
    title = document.add_paragraph(style="Thesis Abstract Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(COVER_TITLE)
    set_east_asia_font(title_run, "黑体", 15, True)

    heading = document.add_paragraph(style="Thesis Abstract Heading")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run("摘  要")
    set_east_asia_font(heading_run, "黑体", 12, True)

    body = document.add_paragraph(style="Thesis Abstract Body")
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing_lines(body, before_lines=0, after_lines=0, line_multiple=1.5)
    body_run = body.add_run(ZH_ABSTRACT)
    set_east_asia_font(body_run, "宋体", 12, False)

    keywords = document.add_paragraph(style="Thesis Abstract Keywords")
    keywords.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing_lines(keywords, before_lines=50, after_lines=0, line_multiple=1.5)
    label_run = keywords.add_run("关键词：")
    set_east_asia_font(label_run, "黑体", 12, True)
    value_run = keywords.add_run(ZH_KEYWORDS)
    set_east_asia_font(value_run, "宋体", 12, False)

    document.add_page_break()

    en_title = document.add_paragraph(style="Thesis Abstract Title")
    en_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    en_title_run = en_title.add_run(EN_TITLE)
    set_east_asia_font(en_title_run, "Times New Roman", 15, True)

    en_heading = document.add_paragraph(style="Thesis Abstract Heading")
    en_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    en_heading_run = en_heading.add_run("Abstract")
    set_east_asia_font(en_heading_run, "Times New Roman", 12, True)

    en_body = document.add_paragraph(style="Thesis Abstract Body")
    en_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing_lines(en_body, before_lines=0, after_lines=0, line_multiple=1.5)
    en_body_run = en_body.add_run(EN_ABSTRACT)
    set_east_asia_font(en_body_run, "Times New Roman", 12, False)

    en_keywords = document.add_paragraph(style="Thesis Abstract Keywords")
    en_keywords.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing_lines(en_keywords, before_lines=50, after_lines=0, line_multiple=1.5)
    en_label_run = en_keywords.add_run("Keywords: ")
    set_east_asia_font(en_label_run, "Times New Roman", 12, True)
    en_value_run = en_keywords.add_run(EN_KEYWORDS)
    set_east_asia_font(en_value_run, "Times New Roman", 12, False)


def add_toc_page(document: Document) -> None:
    title = document.add_paragraph(style="Thesis TOC Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("目  录")
    set_east_asia_font(title_run, "黑体", 14, True)

    toc_entries = load_toc_entries()
    if toc_entries:
        for level, heading, page in toc_entries:
            add_toc_entry(document, level, heading, page)
    else:
        toc_paragraph = document.add_paragraph(style="Thesis TOC 1")
        toc_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing_lines(toc_paragraph, before_lines=0, after_lines=0, line_multiple=1.5)
        add_toc_field(toc_paragraph)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Thesis Body")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing_lines(paragraph, before_lines=0, after_lines=0, line_multiple=1.5)
    add_text_with_citations(paragraph, text, 12)


def add_reference_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Thesis Reference")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing_lines(paragraph, before_lines=0, after_lines=0, line_multiple=1.5)
    paragraph.add_run(text)


def add_text_with_citations(paragraph, text: str, size_pt: int) -> None:
    citation_pattern = re.compile(r"(\[(?:\d+(?:[，,]\d+)*)\])")
    parts = citation_pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        set_east_asia_font(run, "宋体", size_pt, False)
        if citation_pattern.fullmatch(part):
            run.font.superscript = True


def add_figure(document: Document, caption: str, image_path: Path) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.7))

    caption_paragraph = document.add_paragraph(style="Thesis Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing_lines(caption_paragraph, before_lines=0, after_lines=0, line_multiple=1.0)
    caption_paragraph.add_run(caption)


def add_table(document: Document, caption: str, rows: list[list[str]]) -> None:
    if not rows:
        return

    caption_paragraph = document.add_paragraph(style="Thesis Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing_lines(caption_paragraph, before_lines=0, after_lines=0, line_multiple=1.0)
    caption_paragraph.add_run(caption)

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_border(cell, top=row_index == 0, bottom=row_index in {0, len(rows) - 1})
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            set_paragraph_spacing_lines(paragraph, before_lines=0, after_lines=0, line_multiple=1.0)
            run = paragraph.add_run(value.strip())
            set_east_asia_font(run, "宋体", 10, row_index == 0)


def set_cell_border(cell, *, top: bool = False, bottom: bool = False, left: bool = False, right: bool = False) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    enabled_edges = {"top": top, "bottom": bottom, "left": left, "right": right}
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        if enabled_edges[edge]:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "6")
            element.set(qn("w:color"), "000000")
        else:
            element.set(qn("w:val"), "nil")
            element.set(qn("w:sz"), "0")
            element.set(qn("w:color"), "auto")
        element.set(qn("w:space"), "0")


def parse_markdown_table_row(line: str) -> list[str]:
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [cell.strip() for cell in value.split("|")]


def is_markdown_table_separator(line: str) -> bool:
    cells = parse_markdown_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def parse_markdown(text: str) -> list[tuple[str, int | str, str | list[list[str]]]]:
    blocks: list[tuple[str, int | str, str | list[list[str]]]] = []
    paragraph_lines: list[str] = []
    pending_table_caption = ""

    def flush_paragraph() -> None:
        if paragraph_lines:
            blocks.append(("paragraph", "", "".join(paragraph_lines).strip()))
            paragraph_lines.clear()

    lines = text.splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()
        if not line.strip():
            flush_paragraph()
            index += 1
            continue

        if line.strip() == "<!-- PAGE_BREAK -->":
            flush_paragraph()
            blocks.append(("page_break", "", ""))
            index += 1
            continue

        table_caption_match = re.match(r"^(表\d+\.\d+\s+.+)$", line.strip())
        if table_caption_match:
            flush_paragraph()
            pending_table_caption = table_caption_match.group(1)
            index += 1
            continue

        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            flush_paragraph()
            table_rows: list[list[str]] = []
            while index < len(lines):
                table_line = lines[index].strip()
                if not table_line.startswith("|") or "|" not in table_line[1:]:
                    break
                if not is_markdown_table_separator(table_line):
                    table_rows.append(parse_markdown_table_row(table_line))
                index += 1
            if table_rows:
                blocks.append(("table", pending_table_caption, table_rows))
            pending_table_caption = ""
            continue

        figure_match = re.match(r"^!\[(.*?)]\((.*?)\)$", line)
        if figure_match:
            flush_paragraph()
            blocks.append(("figure", figure_match.group(1).strip(), figure_match.group(2).strip()))
            index += 1
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            blocks.append(("heading", level, heading_match.group(2).strip()))
            index += 1
            continue

        paragraph_lines.append(line.strip())
        index += 1

    flush_paragraph()
    return blocks


def load_font(size: int):
    candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill="black") -> None:
    lines = text.split("\n")
    line_heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        line_heights.append(bbox[3] - bbox[1])
    total_height = sum(line_heights) + (len(lines) - 1) * 10
    y = box[1] + ((box[3] - box[1]) - total_height) / 2
    for line, width, height in zip(lines, widths, line_heights):
        x = box[0] + ((box[2] - box[0]) - width) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += height + 10


def wrap_text_to_width(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    wrapped: list[str] = []
    for block in text.split("\n"):
        current = ""
        for ch in block:
            candidate = current + ch
            bbox = draw.textbbox((0, 0), candidate, font=font)
            if current and bbox[2] - bbox[0] > max_width:
                wrapped.append(current)
                current = ch
            else:
                current = candidate
        wrapped.append(current if current else "")
    return wrapped


def fit_box_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, max_size: int, min_size: int = 18):
    inner_width = box[2] - box[0] - 36
    inner_height = box[3] - box[1] - 28
    for size in range(max_size, min_size - 1, -1):
        font = load_font(size)
        lines = wrap_text_to_width(draw, text, font, inner_width)
        bbox = draw.textbbox((0, 0), "测试", font=font)
        line_height = bbox[3] - bbox[1]
        total_height = len(lines) * line_height + max(0, len(lines) - 1) * 8
        max_line_width = max((draw.textbbox((0, 0), line, font=font)[2] for line in lines), default=0)
        if total_height <= inner_height and max_line_width <= inner_width:
            return font, lines
    font = load_font(min_size)
    return font, wrap_text_to_width(draw, text, font, inner_width)


def draw_fitted_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, max_size: int, min_size: int = 18, fill="black") -> None:
    font, lines = fit_box_text(draw, box, text, max_size, min_size)
    joined = "\n".join(lines)
    draw_centered_text(draw, box, joined, font, fill)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], width=4) -> None:
    draw.line([start, end], fill="black", width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - 18 * direction, y2 - 10), (x2 - 18 * direction, y2 + 10)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 10, y2 - 18 * direction), (x2 + 10, y2 - 18 * direction)]
    draw.polygon(points, fill="black")


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font) -> None:
    draw.rounded_rectangle(box, radius=14, outline="black", width=3, fill="#f7f7f7")
    max_size = getattr(font, "size", 28)
    draw_fitted_text(draw, box, text, max_size=max_size, min_size=max(18, max_size - 10))


def save_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42)
    box_font = load_font(28)
    small_font = load_font(24)
    draw_centered_text(draw, (0, 30, 1600, 90), "基于 Spark 的新闻热度监测系统总体架构", title_font)

    boxes = [
        ((80, 210, 300, 360), "数据源\n本地数据集\nRSS/API\n网页正文"),
        ((360, 210, 580, 360), "原始数据层\nJSON/CSV\n采集日志"),
        ((640, 210, 860, 360), "清洗去重层\nHTML 清洗\n时间归一\n强弱去重"),
        ((920, 210, 1140, 360), "Spark 分析层\n热点指数\n趋势统计\n情感/聚类"),
        ((1200, 210, 1420, 360), "后端服务层\nFlask API\n查询筛选\n数据导出"),
    ]
    for box, text in boxes:
        draw_box(draw, box, text, box_font)
    for i in range(len(boxes) - 1):
        arrow(draw, (boxes[i][0][2], 285), (boxes[i + 1][0][0], 285))

    presentation = (500, 560, 1100, 720)
    draw_box(draw, presentation, "前端展示层\nVue3 + ECharts\n可视化大屏 / 新闻查询 / 事件分析 / 服务中心", box_font)
    arrow(draw, (1310, 360), (980, 560))

    draw_centered_text(
        draw,
        (220, 765, 1380, 835),
        "设计思路：先保障数据质量，再完成批处理分析，最后通过接口和页面形成完整结果链路",
        small_font,
    )
    image.save(path)


def save_data_flow_diagram(path: Path) -> None:
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42)
    box_font = load_font(28)
    draw_centered_text(draw, (0, 30, 1600, 90), "新闻数据处理流程", title_font)

    rows = [
        ((120, 150, 430, 260), "采集入口\n本地导入 / RSS / 网页解析"),
        ((645, 150, 955, 260), "字段标准化\n统一 news_id、title、\ncontent 等字段"),
        ((1170, 150, 1480, 260), "原始数据存储\n保存原始新闻与采集记录"),
        ((1170, 390, 1480, 500), "清洗与去重\n噪声清理、时间标准化、强弱去重"),
        ((645, 390, 955, 500), "Spark 批处理\n热度、趋势、情感、事件聚类"),
        ((120, 390, 430, 500), "结果落库\nParquet / JSON / SQLite"),
        ((120, 640, 430, 750), "Flask 接口\n查询、筛选、导出"),
        ((645, 640, 955, 750), "前端展示\n大屏、检索、事件分析"),
    ]
    for box, text in rows:
        draw_box(draw, box, text, box_font)

    arrows = [
        ((430, 205), (645, 205)),
        ((955, 205), (1170, 205)),
        ((1325, 260), (1325, 390)),
        ((1170, 445), (955, 445)),
        ((645, 445), (430, 445)),
        ((275, 500), (275, 640)),
        ((430, 695), (645, 695)),
    ]
    for start, end in arrows:
        arrow(draw, start, end)
    image.save(path)


def save_module_diagram(path: Path) -> None:
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42)
    box_font = load_font(28)
    draw_centered_text(draw, (0, 30, 1600, 90), "系统功能模块设计", title_font)

    center = (610, 365, 990, 515)
    draw_box(draw, center, "新闻热度监测系统", box_font)
    modules = [
        ((110, 160, 430, 280), "新闻采集模块\n数据集、RSS、网页解析"),
        ((640, 160, 960, 280), "清洗去重模块\n文本清洗、格式统一、相似去重"),
        ((1170, 160, 1490, 280), "热点分析模块\n热度指数、榜单统计"),
        ((110, 620, 430, 740), "趋势情感模块\n关键词趋势、情感概览"),
        ((640, 620, 960, 740), "事件聚类模块\nTF-IDF、相似事件归并"),
        ((1170, 620, 1490, 740), "查询导出模块\n筛选、检索、CSV/Excel 导出"),
    ]
    for box, text in modules:
        draw_box(draw, box, text, box_font)
        arrow(draw, ((box[0] + box[2]) // 2, box[3] if box[1] < center[1] else box[1]), ((center[0] + center[2]) // 2, center[1] if box[1] < center[1] else center[3]))
    image.save(path)


def crop_image(source_path: Path, target_path: Path, box: tuple[int, int, int, int]) -> None:
    if not source_path.exists():
        return
    with Image.open(source_path) as image:
        left, top, right, bottom = box
        left = max(0, min(left, image.width - 1))
        top = max(0, min(top, image.height - 1))
        right = max(left + 1, min(right, image.width))
        bottom = max(top + 1, min(bottom, image.height))
        cropped = image.crop((left, top, right, bottom))
        cropped.save(target_path)


def ensure_ui_figures() -> None:
    crop_specs = [
        ("dashboard-page.png", "dashboard_metrics.png", (0, 0, 2561, 760)),
        ("dashboard-page.png", "dashboard_analysis.png", (330, 720, 2561, 1950)),
        ("news-page.png", "news_filters.png", (0, 0, 2561, 560)),
        ("news-page.png", "news_results.png", (330, 380, 2561, 1480)),
        ("event-page.png", "events_clusters.png", (0, 0, 2561, 1220)),
        ("event-page.png", "events_detail.png", (330, 1860, 2561, 2860)),
        ("event-page.png", "events_timeline.png", (330, 2060, 2561, 3142)),
        ("service-page.png", "services_overview.png", (0, 0, 1781, 1650)),
    ]
    for source_name, target_name, box in crop_specs:
        crop_image(ROOT / "output/doc" / source_name, FIGURE_DIR / target_name, box)


def ensure_generated_figures() -> None:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    save_architecture_diagram(FIGURE_DIR / "system_architecture.png")
    save_data_flow_diagram(FIGURE_DIR / "data_flow.png")
    save_module_diagram(FIGURE_DIR / "function_modules.png")
    ensure_ui_figures()


def build() -> Path:
    ensure_generated_figures()
    document = Document(str(TEMPLATE))
    clear_document_body(document)
    configure_styles(document)
    enable_update_fields(document)

    configure_cover_section(document.sections[0])
    add_cover(document)

    abstract_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    configure_front_matter_section(abstract_section)
    add_abstract_pages(document)

    toc_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    configure_running_section(toc_section, "upperRoman", 1)
    add_toc_page(document)

    body_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    configure_running_section(body_section, "decimal", 1)

    blocks = parse_markdown(SOURCE.read_text(encoding="utf-8"))
    first_h1 = True
    in_references = False
    for kind, meta, text in blocks:
        if kind == "heading":
            add_heading(document, text, int(meta), first_h1)
            if int(meta) == 1 and first_h1:
                first_h1 = False
            in_references = text == "参考文献"
            if text == "致谢":
                in_references = False
        elif kind == "figure":
            image_path = ROOT / text
            add_figure(document, str(meta), image_path)
        elif kind == "table":
            add_table(document, str(meta), text)  # type: ignore[arg-type]
        elif kind == "page_break":
            document.add_page_break()
        else:
            if in_references:
                add_reference_paragraph(document, text)
            else:
                add_paragraph(document, text)

    document.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    out = build()
    print(out)

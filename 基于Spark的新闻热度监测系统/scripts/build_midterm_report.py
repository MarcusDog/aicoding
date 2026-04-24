from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TMP_DIR = ROOT.parent / "tmp" / "docs"
TEMPLATE_PATH = TMP_DIR / "midterm_form_converted.docx"
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_PATH = OUTPUT_DIR / "基于Spark的新闻热度监测系统中期报告.docx"

TITLE = "基于Spark的新闻热度监测系统的设计与实现"

WORK_TEXT = (
    "本人在该毕业设计中主要负责新闻热度监测系统的需求分析、总体方案设计、数据采集与处理链路实现、"
    "后端接口开发、前端展示页面开发以及阶段性测试与文档整理。具体包括：第一，围绕新闻热度监测场景完成系统需求梳理，"
    "明确数据来源、处理流程、分析指标和展示方式；第二，设计新闻统一数据结构，完成原始数据、清洗结果、分析结果和导出结果的存储方案；"
    "第三，基于 Python 与 Spark 相关技术实现新闻采集、批量导入、清洗去重、热度计算、关键词趋势分析、情感分析、事件聚类和预警分析等核心功能；"
    "第四，基于 Flask 构建后端接口，支持大屏展示、新闻查询、事件详情分析、传播力概览和数据导出；第五，基于 Vue3 和 ECharts 完成系统前端页面，"
    "实现热点展示、趋势图表、事件分析和服务简报等可视化功能；第六，对系统进行联调测试、结果校验和中期阶段文档整理，为后续系统优化、论文撰写和毕业答辩做好准备。"
)

PROGRESS_BLOCKS = [
    (
        "1. 项目完成进度：约 65%",
        [
            "从当前项目进展来看，系统已经完成中期检查所要求的核心闭环，整体进度约为 65%。目前已实现“数据采集 - 数据清洗与去重 - 热度分析 - 结果存储 - 接口服务 - 前端展示”的基本链路，项目已经具备可运行、可演示、可继续扩展的原型系统基础。",
        ],
    ),
    (
        "2. 已完成的前期调研与方案设计工作",
        [
            "项目启动后，首先结合任务书要求对新闻热度监测系统的应用场景、实现目标和技术路线进行了梳理，明确中期阶段以“先跑通系统闭环，再逐步增强分析能力”为主要实施思路。围绕这一思路，完成了项目总规划文档、实施方案文档、数据源接入矩阵和阶段性优化清单，明确了系统分层结构、功能边界和中后期优化方向，为后续编码实现提供了比较清晰的路线。",
        ],
    ),
    (
        "3. 已完成的数据采集与数据底座建设",
        [
            "在数据采集方面，已经完成多来源采集方案的基础实现，支持样例数据导入、批量数据集导入、RSS 聚合采集、网页正文抓取和开放热点 API 接入。当前系统已整理出统一的数据源目录配置文件，能够对人民网、新华网、BBC、Reuters、腾讯新闻、网易新闻、36Kr、虎嗅以及微博热点、知乎热点、Hacker News、Lobsters 等来源进行分类管理。",
            "系统支持 csv、json、jsonl 等多种离线数据导入形式，并可将原始采集结果按批次写入 latest_raw.json 及历史归档目录。根据当前已有运行结果，原始新闻数据规模已达到 12490 条，覆盖 29 个新闻来源，为中期阶段的分析验证提供了数据基础。",
        ],
    ),
    (
        "4. 已完成的数据清洗、去重与存储工作",
        [
            "围绕新闻数据质量问题，已完成基础清洗与去重模块的实现。系统能够对原始新闻进行 HTML 标签清洗、空白字符规范化、时间字段标准化、关键词抽取以及结构化字段补全，并通过 URL、标题哈希和 SimHash 相结合的方式执行强去重与弱去重，降低转载内容对热点结果的干扰。",
            "在数据存储方面，系统已经形成原始数据、清洗结果、分析结果多层保存机制，能够将处理结果输出为 JSON、Parquet 以及数据库记录，方便后续统计分析、页面展示和论文撰写时的数据取证。",
        ],
    ),
    (
        "5. 已完成的热度分析与核心算法模块",
        [
            "在分析层面，系统已实现基于新闻数量、来源权重、情感强度、互动指标和时效衰减的可解释热度计算方法，并支持常规场景、危机场景、政策场景三种热度评分口径。与此同时，系统还实现了关键词趋势统计、基础情感分析、事件聚类和规则预警等模块。",
            "其中，情感分析目前采用词典规则法完成正面、中性、负面分类，事件聚类采用 TF-IDF 与层次聚类生成热点事件簇，预警模块能够识别新闻量突增和负面占比偏高等情况。根据现有处理结果，系统已经生成 11876 条清洗结果、20 条热点主题、45 条关键词趋势、8 个事件簇和 19 条预警信息，说明中期阶段的分析流程已经跑通。",
        ],
    ),
    (
        "6. 已完成的后端服务与接口开发",
        [
            "后端部分已基于 Flask 完成接口封装，当前已经提供新闻概览、新闻检索、热点列表、关键词趋势、情感汇总、事件簇查询、事件详情、预警简报、服务简报、传播力概览以及新闻导出等接口。事件详情接口已能输出时间线、传播路径和 5W1H 结构化摘要，导出接口支持按查询结果导出为 CSV 或 Excel 文件。",
            "通过这些接口，系统已经能够支撑前端页面展示和中期阶段的功能演示。",
        ],
    ),
    (
        "7. 已完成的前端页面与系统展示工作",
        [
            "前端部分已基于 Vue3、Vite 和 ECharts 完成多个核心页面开发，现已具备新闻监测可视化大屏、新闻查询页、事件分析页、传播力看板和服务简报页等功能页面。大屏页面能够展示新闻总量、来源覆盖、事件簇数量、平均热度、来源分布、情感分布、关键词趋势和预警列表，并支持触发数据采集与分析。",
            "查询页和分析页已能配合后端接口完成数据浏览与分析展示，说明系统在中期阶段已经具备基本的可视化演示能力。",
        ],
    ),
    (
        "8. 已完成的测试与阶段性成果整理",
        [
            "为保证系统的可运行性，当前后端已编写基础自动化测试，覆盖数据清洗去重、分析输出、批量导入、来源目录加载、接口可访问性、事件详情、预警简报和传播力概览等关键功能。当前测试结果显示，后端共 14 项测试全部通过，说明中期阶段核心链路具备一定稳定性。",
            "除此之外，项目还完成了实施文档、答辩准备笔记和优化路线文档的整理，为后续继续完善系统和撰写毕业论文提供了较好的基础。",
        ],
    ),
]

PLAN_BLOCKS = [
    (
        "1. 目前存在的主要问题",
        [
            "虽然系统主体链路已经跑通，但中期阶段仍存在一些比较明显的问题。首先，Spark 分析链路目前采用“优先 Spark、异常时退化为 Pandas”的实现方式，本地 Spark 运行环境和更大规模数据下的性能验证还不够充分；其次，情感分析和事件聚类目前仍以基础算法为主，模型精度和中文语义识别能力还有较大提升空间。",
            "再次，网页正文抓取对不同站点结构的适配能力有限，部分来源的稳定性和正文抽取质量仍需加强；另外，前端页面虽然已经完成核心展示，但筛选联动、空状态处理、交互细节和截图友好性仍需优化；最后，系统测试目前主要集中在后端基础功能，对真实新闻样本、接口压力、前端页面联调以及论文所需实验数据支撑还不够完整。",
        ],
    ),
    (
        "2. 下一步的主要设计任务",
        [
            "下一阶段将重点围绕“增强数据质量、提升分析效果、补齐测试材料、完善论文支撑”四个方向展开。第一，继续扩展稳定新闻源，完善来源管理、采集日志、失败重试和批次追踪机制，提高数据采集的可控性；第二，优化热度指标权重和关键词处理规则，补充同义词合并、来源权威度映射和场景化评分策略。",
            "第三，尝试引入更适合中文新闻场景的情感分析模型和句向量方法，提升情感分析与事件聚类的准确性；第四，进一步完善后端参数校验、接口联调和导出功能，并优化前端页面筛选条件、状态反馈和图表联动效果；第五，补充系统测试、运行截图、实验统计结果和论文素材，为后续毕业设计说明书撰写和答辩演示提供完整证据。",
        ],
    ),
    (
        "3. 具体设想与时间安排",
        [
            "在后续安排上，计划先完成数据采集和分析模块的稳定性优化，再完善前后端联调和页面表现，最后集中整理测试材料和论文支撑内容。具体而言，近期将优先完成真实新闻源补充、采集策略优化和分析参数调试；随后完成情感分析、事件聚类和传播力展示的增强；在系统功能基本稳定后，再系统整理测试结果、系统截图、图表和文档，逐步完成毕业论文中系统设计、实现和测试章节所需材料。",
            "通过上述安排，确保项目从“中期可运行原型”顺利推进到“后期可答辩、可交付系统”。",
        ],
    ),
]


def set_run_font(run, size: float = 12.0, bold: bool = False) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(paragraph, *, first_line_indent: bool) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(24) if first_line_indent else Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def clear_cell_body(cell) -> None:
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    cell.paragraphs[0].clear()


def trim_leading_empty_paragraphs(cell) -> None:
    while len(cell.paragraphs) > 1 and not cell.paragraphs[0].text.strip():
        p = cell.paragraphs[0]._element
        p.getparent().remove(p)


def write_first_paragraph(cell, text: str, *, indent: bool, heading: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    set_paragraph_format(paragraph, first_line_indent=indent)
    run = paragraph.add_run(text)
    set_run_font(run, size=12.0, bold=False)


def append_body_paragraph(cell, text: str, *, indent: bool = True) -> None:
    paragraph = cell.add_paragraph()
    set_paragraph_format(paragraph, first_line_indent=indent)
    run = paragraph.add_run(text)
    set_run_font(run, size=12.0, bold=False)


def append_heading_paragraph(cell, text: str) -> None:
    paragraph = cell.add_paragraph()
    set_paragraph_format(paragraph, first_line_indent=False)
    run = paragraph.add_run(text)
    set_run_font(run, size=12.0, bold=False)


def set_cell_single_paragraph(cell, text: str, *, center: bool = False) -> None:
    clear_cell_body(cell)
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=12.0, bold=False)


def fill_metadata(table) -> None:
    set_cell_single_paragraph(table.cell(0, 1), "软件工程")
    set_cell_single_paragraph(table.cell(0, 4), "")
    set_cell_single_paragraph(table.cell(0, 6), "")
    set_cell_single_paragraph(table.cell(1, 1), "")
    set_cell_single_paragraph(table.cell(1, 3), "")
    set_cell_single_paragraph(table.cell(1, 6), "")
    set_cell_single_paragraph(table.cell(2, 1), TITLE)


def fill_work_cell(table) -> None:
    cell = table.cell(3, 4)
    clear_cell_body(cell)
    write_first_paragraph(cell, WORK_TEXT, indent=True)
    trim_leading_empty_paragraphs(cell)


def fill_progress_cell(table) -> None:
    cell = table.cell(4, 4)
    clear_cell_body(cell)
    first_heading, first_paragraphs = PROGRESS_BLOCKS[0]
    write_first_paragraph(cell, first_heading, indent=False)
    for text in first_paragraphs:
        append_body_paragraph(cell, text, indent=True)
    for heading, paragraphs in PROGRESS_BLOCKS[1:]:
        append_heading_paragraph(cell, heading)
        for text in paragraphs:
            append_body_paragraph(cell, text, indent=True)
    trim_leading_empty_paragraphs(cell)


def fill_plan_cell(table) -> None:
    cell = table.cell(5, 4)
    clear_cell_body(cell)
    first_heading, first_paragraphs = PLAN_BLOCKS[0]
    write_first_paragraph(cell, first_heading, indent=False)
    for text in first_paragraphs:
        append_body_paragraph(cell, text, indent=True)
    for heading, paragraphs in PLAN_BLOCKS[1:]:
        append_heading_paragraph(cell, heading)
        for text in paragraphs:
            append_body_paragraph(cell, text, indent=True)
    trim_leading_empty_paragraphs(cell)


def clean_teacher_cell(table) -> None:
    cell = table.cell(6, 4)
    clear_cell_body(cell)
    write_first_paragraph(cell, "指导教师对该学生前期设计工作的评价（是否同意继续设计工作）", indent=False)
    append_body_paragraph(cell, "指导教师签字：                        年    月    日", indent=False)
    trim_leading_empty_paragraphs(cell)


def build_report() -> Path:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"template not found: {TEMPLATE_PATH}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE_PATH))
    table = doc.tables[0]

    fill_metadata(table)
    fill_work_cell(table)
    fill_progress_cell(table)
    fill_plan_cell(table)
    clean_teacher_cell(table)

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_report()
    print(json.dumps({"output": str(output)}, ensure_ascii=False))
